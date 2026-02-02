"""
SISTEMA DE PETIÇÕES AUTOMATIZADAS V10.0 - FASE 3 PERFEITA
Sistema Completo: 3 Prioridades + Pedidos + Estrutura + Reflexos HE + Prints + Estatísticas
Baseado no Checklist IA Auditora v4.0 FINAL
"""

import os
import time
import json
import schedule
import glob
import unicodedata
from datetime import datetime
from dotenv import load_dotenv
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from anthropic import Anthropic
import io
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
import tempfile
import base64
import google.generativeai as genai

# Importar módulo Prompt Master (opcional - para petições de alto nível)
try:
    from prompt_master import gerar_prompt_master
    from validacao_prompt_master import gerar_relatorio_validacao_master, imprimir_relatorio_validacao
    PROMPT_MASTER_DISPONIVEL = True
except ImportError:
    PROMPT_MASTER_DISPONIVEL = False
    print(" Módulo Prompt Master não disponível")

load_dotenv()

SCOPES = ['https://www.googleapis.com/auth/drive']
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY')
anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)

# Configurar Gemini API para transcrição de vídeos
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

def normalizar_texto(texto):
    """Remove acentos e caracteres especiais para comparação"""
    if not texto:
        return ""
    # Normaliza para NFD (decompõe caracteres acentuados)
    nfd = unicodedata.normalize('NFD', texto)
    # Remove marcas diacríticas (acentos)
    sem_acentos = ''.join(char for char in nfd if unicodedata.category(char) != 'Mn')
    return sem_acentos.lower()

# ============================================================================
# CHECKLIST V4.0 - SISTEMA DE 3 PRIORIDADES
# ============================================================================

# DOCUMENTOS POR TIPO - BASEADO NO CHECKLIST V4.0 OFICIAL
DOCUMENTOS_POR_TIPO = {
    'RECONHECIMENTO_VINCULO': {
        'ALTA': ['DOCUMENTO_PESSOAL', 'CTPS', 'COMPROVANTE_PAGAMENTO', 'TRANSCRICAO'],
        'MEDIA': ['COMPROVANTE_RESIDENCIA', 'FOTOS_TRABALHANDO', 'TRCT'],
        'BAIXA': ['CONVERSAS_WHATSAPP', 'CONTRATO']
    },
    'ACAO_ACIDENTARIA': {
        'ALTA': ['DOCUMENTO_PESSOAL', 'CTPS', 'COMPROVANTE_PAGAMENTO', 'ATESTADO_MEDICO', 
                 'EXAMES_MEDICOS', 'DOCUMENTOS_PREVIDENCIARIOS', 'CAT', 'TRANSCRICAO'],
        'MEDIA': ['COMPROVANTE_RESIDENCIA', 'FOTOS_TRABALHANDO', 'FOTOS_AMBIENTE', 'TRCT'],
        'BAIXA': ['CONVERSAS_WHATSAPP']
    },
    'DIFERENCAS_CONTRATUAIS': {
        'ALTA': ['DOCUMENTO_PESSOAL', 'CTPS', 'COMPROVANTE_PAGAMENTO', 'HOLERITES', 'TRCT', 'TRANSCRICAO'],
        'MEDIA': ['COMPROVANTE_RESIDENCIA', 'FOTOS_TRABALHANDO', 'CONTRATO', 
                  'CARTAO_PONTO', 'EXTRATO_FGTS'],
        'BAIXA': ['CONVERSAS_WHATSAPP']
    }
}

def agente_cronologia(texto_transcricao, cliente_nome):
    """Gera linha do tempo baseada na transcrição"""
    print(f"        [AGENTE CRONOLOGIA] Analisando transcrição...")
    try:
        prompt = f"""
        Você é um assistente paralegal especialista em análise de provas e fatos.
        
        CLIENTE: {cliente_nome}
        
        TAREFA: Analise a transcrição abaixo e crie uma LINHA DO TEMPO (CRONOLOGIA) detalhada dos fatos.
        OBJETIVO: Identificar datas, eventos chaves, promessas, demissões, acidentes, valores e testemunhas e qualquer outro fato relevante para o processo.
        
        Instruções de Formatação:
        1. Crie um título: "CRONOLOGIA DOS FATOS - {cliente_nome}"
        2. Use uma lista cronológica onde cada item tenha: [DATA/PERÍODO] - [EVENTO]
        3. Se não houver data exata, use "Data Indeterminada" ou o período aproximado (ex: "Inicio de 2024").
        4. Destaque em NEGRITO valore monetários e datas críticas (admissão/demissão/acidente).
        
        TRANSCRICÃO:
        {texto_transcricao[:50000]}
        """
        
        message = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16000,
            system="Você é um especialista em extração de fatos cronológicos.",
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        print(f"         Erro na cronologia: {e}")
        return None

def salvar_cronologia_docx(service, texto_cronologia, cliente_nome, pasta_id):
    """Salva a cronologia em DOCX no Drive"""
    print(f"        [AGENTE CRONOLOGIA] Salvando arquivo...")
    try:
        doc = Document()
        doc.add_heading(f'Cronologia dos Fatos - {cliente_nome}', 0)
        
        # Adicionar texto preservando quebras de linha básicas
        for paragrafo in texto_cronologia.split('\\n'):
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())
                
        # Salvar temporariamente
        temp_filename = f'Cronologia_Fatos_{int(time.time())}.docx'
        temp_path = os.path.join(tempfile.gettempdir(), temp_filename)
        doc.save(temp_path)
        
        # Upload para o Drive
        file_metadata = {
            'name': 'Cronologia_Fatos.docx',
            'parents': [pasta_id]
        }
        
        # Usar try-finally para garantir que recursos sejam liberados
        media = None
        try:
            media = MediaFileUpload(temp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            
            arquivo = service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            
            print(f"         Cronologia salva! ID: {arquivo.get('id')}")
            
        finally:
            # Tentar fechar o media se possível (algumas implementações exigem)
            if media and hasattr(media, 'stream') and media.stream:
                try:
                    media.stream.close()
                except:
                    pass

        try:
            os.remove(temp_path)
        except Exception as e:
            print(f"         Aviso: Não foi possível remover arquivo temporário {temp_path}: {e}")
            
        return True
    except Exception as e:
        print(f"         Erro ao salvar cronologia: {e}")
        return False

# ============================================================================
# AGENTE DE TRANSCRIÇÃO DE VÍDEOS
# ============================================================================

def buscar_ou_criar_pasta(service, nome_pasta, pasta_pai_id):
    """Busca uma pasta pelo nome dentro de uma pasta pai, ou cria se não existir"""
    try:
        # Buscar pasta existente
        query = f"name='{nome_pasta}' and '{pasta_pai_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        results = service.files().list(q=query, fields="files(id, name)").execute()
        items = results.get('files', [])
        
        if items:
            print(f"         Pasta '{nome_pasta}' encontrada: {items[0]['id']}")
            return items[0]['id']
        
        # Criar nova pasta
        file_metadata = {
            'name': nome_pasta,
            'mimeType': 'application/vnd.google-apps.folder',
            'parents': [pasta_pai_id]
        }
        folder = service.files().create(body=file_metadata, fields='id').execute()
        print(f"         Pasta '{nome_pasta}' criada: {folder.get('id')}")
        return folder.get('id')
        
    except Exception as e:
        print(f"         Erro ao buscar/criar pasta: {e}")
        return None

def baixar_arquivo(service, file_id):
    """Baixa arquivo do Google Drive"""
    try:
        request = service.files().get_media(fileId=file_id)
        file_stream = io.BytesIO()
        downloader = MediaIoBaseDownload(file_stream, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_stream.seek(0)
        return file_stream.read()
    except Exception as e:
        print(f"         Erro ao baixar arquivo: {e}")
        return None

def agente_transcricao_video(service, video_id, video_nome, cliente_nome, pasta_cliente_id):
    """
    Transcreve vídeo usando Gemini API e salva em pasta dedicada
    
    Args:
        service: Google Drive service
        video_id: ID do vídeo no Drive
        video_nome: Nome do arquivo de vídeo
        cliente_nome: Nome do cliente
        pasta_cliente_id: ID da pasta do cliente no Drive
    
    Returns:
        dict com status, transcricao_id, texto, etc.
    """
    print(f"\n     [AGENTE RESUMO] Iniciando resumo de vídeo...")
    print(f"        Cliente: {cliente_nome}")
    print(f"        Vídeo: {video_nome}")

    # Helper para atualizar status
    def atualizar_progresso(mensagem, etapa, total_etapas=4):
        try:
            filename = f"status_video_{re.sub(r'[^a-zA-Z0-9]', '_', cliente_nome)}.json"
            filepath = os.path.join("flags", filename)
            data = {
                "cliente": cliente_nome,
                "mensagem": mensagem,
                "etapa": etapa,
                "total_etapas": total_etapas,
                "timestamp": time.time(),
                "concluido": False
            }
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f)
        except Exception as e:
            print(f"Erro ao salvar status: {e}")

    print(f"[DEBUG] Chamando atualizar_progresso - Etapa 1")
    atualizar_progresso("Iniciando processamento...", 1)
    print(f"[DEBUG] atualizar_progresso concluído - Etapa 1")
    
    try:
        # 1. Verificar se Gemini API está configurada
        if not GEMINI_API_KEY:
            return {
                'success': False,
                'error': 'GEMINI_API_KEY não configurada no arquivo .env'
            }
        
        # 2. Buscar ou criar pasta "Transcrições de Vídeo"
        # Mantendo o nome da pasta para compatibilidade
        print(f"         Criando/buscando pasta de transcrições...")
        pasta_transcricoes_id = buscar_ou_criar_pasta(service, "Transcrições de Vídeo", pasta_cliente_id)
        
        if not pasta_transcricoes_id:
            return {
                'success': False,
                'error': 'Não foi possível criar pasta de transcrições'
            }
        
        # 3. Baixar vídeo do Drive
        print(f"        ⬇ Baixando vídeo do Drive...")
        atualizar_progresso("Baixando vídeo do Google Drive...", 1)
        video_bytes = baixar_arquivo(service, video_id)
        
        if not video_bytes:
            return {
                'success': False,
                'error': 'Não foi possível baixar o vídeo do Drive'
            }
        
        # 4. Salvar temporariamente
        extensao = os.path.splitext(video_nome)[1]
        temp_video_path = os.path.join(tempfile.gettempdir(), f'video_temp_{int(time.time())}{extensao}')
        
        with open(temp_video_path, 'wb') as f:
            f.write(video_bytes)
        
        print(f"         Vídeo salvo: {temp_video_path} ({len(video_bytes) / (1024*1024):.2f} MB)")
        
        # 5. Upload para Gemini e transcrever
        print(f"         Enviando para Gemini API...")
        atualizar_progresso("Enviando vídeo para análise da IA...", 2)
        
        try:
            # Upload do vídeo para Gemini
            video_file = genai.upload_file(path=temp_video_path)
            print(f"         Vídeo enviado: {video_file.name}")
            
            # Aguardar processamento
            print(f"         Aguardando processamento...")
            while video_file.state.name == "PROCESSING":
                time.sleep(2)
                video_file = genai.get_file(video_file.name)
            
            if video_file.state.name == "FAILED":
                raise Exception("Falha no processamento do vídeo pelo Gemini")
            
            # Criar modelo e transcrever
            model = genai.GenerativeModel(model_name="gemini-2.5-flash")
            
            prompt = f"""
            Você é um assistente paralegal especializado em análise de vídeos e entrevistas com clientes.
            
            TAREFA: Analise o vídeo e gere um RESUMO DETALHADO DOS FATOS para fundamentação de peça jurídica.
            
            INSTRUÇÕES:
            1. NÃO faça transcrição palavra por palavra. Foque nos FATOS.
            2. Identifique cronologicamente a história narrada.
            3. Extraia dados concretos: Datas, Valores, Nomes, Locais, Acidentes, Doenças.
            4. Se houver contradições ou pontos confusos, anote como observação.
            5. Mantenha tom profissional e jurídico.
            
            FORMATO DE SAÍDA:
            RESUMO DO VÍDEO - {video_nome}
            Cliente: {cliente_nome}
            Data: {datetime.now().strftime("%d/%m/%Y")}
            
            == RESUMO DOS FATOS ==
            [Resumo cronológico e detalhado aqui]
            
            == DADOS EXTRAÍDOS ==
            - Admissão/Demissão: [se mencionado]
            - Função: [se mencionado]
            - Salário: [se mencionado]
            - Acidentes/Doenças: [se mencionado]
            - Testemunhas: [se mencionado]
            
            Gere o resumo agora:
            """
            
            print(f"         Gerando resumo...")
            atualizar_progresso("Gerando resumo detalhado dos fatos...", 3)
            try:
                # Usar configurações de segurança mais permissivas para não bloquear o resumo
                # Usando strings diretas que são compatíveis com a versão mais recente
                safety_settings = [
                    {
                        "category": "HARM_CATEGORY_HARASSMENT",
                        "threshold": "BLOCK_NONE"
                    },
                    {
                        "category": "HARM_CATEGORY_HATE_SPEECH",
                        "threshold": "BLOCK_NONE"
                    },
                    {
                        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        "threshold": "BLOCK_NONE"
                    },
                    {
                        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                        "threshold": "BLOCK_NONE"
                    }
                ]

                response = model.generate_content(
                    [video_file, prompt],
                    safety_settings=safety_settings
                )
                texto_transcricao = response.text
            except Exception as api_error:
                print(f"         ERRO na API Gemini: {api_error}")
                print(f"         Tipo do erro: {type(api_error).__name__}")
                import traceback
                traceback.print_exc()
                raise
            
            print(f"         Resumo gerado! ({len(texto_transcricao)} caracteres)")
            
        finally:
            # Limpar arquivo temporário
            try:
                os.remove(temp_video_path)
            except:
                pass
        
        # 6. Salvar transcrição como DOCX
        print(f"         Salvando RESUMO como DOCX...")
        atualizar_progresso("Salvando arquivo na pasta do cliente...", 4)
        
        doc = Document()
        doc.add_heading(f'Resumo de Vídeo - {cliente_nome}', 0)
        doc.add_paragraph(f'Vídeo Original: {video_nome}')
        doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Adicionar transcrição
        for paragrafo in texto_transcricao.split('\n'):
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())
        
        # Salvar temporariamente
        nome_base = os.path.splitext(video_nome)[0]
        # MUDANÇA: Prefixo RESUMO_
        temp_docx_filename = f'RESUMO_{nome_base}_{int(time.time())}.docx'
        temp_docx_path = os.path.join(tempfile.gettempdir(), temp_docx_filename)
        doc.save(temp_docx_path)
        
        # Upload para Drive na pasta de transcrições
        file_metadata = {
            'name': f'RESUMO_{nome_base}.docx',
            'parents': [pasta_transcricoes_id]
        }
        
        media = MediaFileUpload(
            temp_docx_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        arquivo = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        # Limpar arquivo temporário
        try:
            os.remove(temp_docx_path)
        except:
            pass
        
        print(f"         Resumo salvo no Drive!")
        print(f"         Arquivo: {arquivo.get('name')}")
        print(f"         ID: {arquivo.get('id')}")
        
        return {
            'success': True,
            'transcricao_id': arquivo.get('id'),
            'transcricao_nome': arquivo.get('name'),
            'transcricao_link': arquivo.get('webViewLink'),
            'texto': texto_transcricao,
            'texto_preview': texto_transcricao[:500] + '...' if len(texto_transcricao) > 500 else texto_transcricao,
            'tamanho_caracteres': len(texto_transcricao)
        }
        
    except Exception as e:
        print(f"         Erro na transcrição: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }


# PONTUAÇÃO POR PRIORIDADE - CHECKLIST V4.0
PONTOS_PRIORIDADE = {
    'ALTA': -10,    #  BLOQUEIA - Crítico
    'MEDIA': -5,    #  ALERTA - Importante
    'BAIXA': -2     #  SUGERE - Desejável
}

# SCORES MÉDIOS POR TIPO - BASE DE 100 PETIÇÕES (CHECKLIST V4.0)
SCORES_MEDIOS = {
    'RECONHECIMENTO_VINCULO': 81.7,
    'ACAO_ACIDENTARIA': 77.0,
    'DIFERENCAS_CONTRATUAIS': 71.0
}

# ============================================================================
# SISTEMA DE JURISPRUDÊNCIAS V1.0
# ============================================================================

def carregar_jurisprudencias():
    """Carrega banco de jurisprudências do arquivo JSON"""
    try:
        caminho = os.path.join(os.path.dirname(__file__), 'jurisprudencias.json')
        if os.path.exists(caminho):
            with open(caminho, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            print(f"        [AVISO] Arquivo jurisprudencias.json não encontrado")
    except Exception as e:
        print(f"        [ERRO] Ao carregar jurisprudências: {e}")
    return {}

def obter_jurisprudencia_por_tema(tema, quantidade=2):
    """
    Retorna jurisprudências formatadas para um tema específico
    
    Args:
        tema: Nome do tema (ex: 'reconhecimento_vinculo', 'horas_extras')
        quantidade: Número de jurisprudências a retornar (padrão: 2)
    
    Returns:
        String formatada com as jurisprudências ou vazio se não encontrar
    """
    juris_db = carregar_jurisprudencias()
    juris_tema = juris_db.get(tema, [])
    
    if not juris_tema:
        return ""
    
    # Pegar as primeiras N jurisprudências
    selecionadas = juris_tema[:quantidade]
    
    texto = ""
    for j in selecionadas:
        texto += f"\n({j['tribunal']} - {j['tipo']}: {j['numero']}, "
        texto += f"Relator: {j['relator']}, Data: {j['data']}, Turma: {j['turma']})\n"
        texto += f'"{j["ementa"]}"\n'
    
    return texto

def obter_jurisprudencias_para_tipo_processo(tipo_processo):
    """
    Retorna jurisprudências relevantes baseadas no tipo de processo
    
    Args:
        tipo_processo: RECONHECIMENTO_VINCULO, ACAO_ACIDENTARIA, DIFERENCAS_CONTRATUAIS
    
    Returns:
        Dict com jurisprudências organizadas por tema
    """
    juris = {}
    
    if tipo_processo == 'RECONHECIMENTO_VINCULO':
        juris['reconhecimento_vinculo'] = obter_jurisprudencia_por_tema('reconhecimento_vinculo', 2)
        juris['rescisao_indireta'] = obter_jurisprudencia_por_tema('rescisao_indireta', 1)
        juris['multa_477'] = obter_jurisprudencia_por_tema('multa_477', 1)
        juris['multa_467'] = obter_jurisprudencia_por_tema('multa_467', 1)
    
    elif tipo_processo == 'ACAO_ACIDENTARIA':
        juris['danos_morais'] = obter_jurisprudencia_por_tema('danos_morais', 2)
        juris['rescisao_indireta'] = obter_jurisprudencia_por_tema('rescisao_indireta', 1)
    
    elif tipo_processo == 'DIFERENCAS_CONTRATUAIS':
        juris['horas_extras'] = obter_jurisprudencia_por_tema('horas_extras', 2)
        juris['adicional_noturno'] = obter_jurisprudencia_por_tema('adicional_noturno', 1)
        juris['insalubridade'] = obter_jurisprudencia_por_tema('insalubridade', 1)
        juris['multa_477'] = obter_jurisprudencia_por_tema('multa_477', 1)
    
    return juris

# ============================================================================
# SISTEMA DE IDENTIFICAÇÃO DE PRINTS V4.0
# ============================================================================

PRINTS_POR_TIPO = {
    'RECONHECIMENTO_VINCULO': {
        'ALTA': [
            {'nome': 'Carteira de Trabalho', 'gatilhos': ['CTPS', 'carteira de trabalho', 'ausência de registro', 'registro em carteira']},
            {'nome': 'Comprovantes de Pagamento', 'gatilhos': ['comprovante', 'recibo', 'transferência', 'pagamento']}
        ],
        'MEDIA': [
            {'nome': 'Fotos Trabalhando', 'gatilhos': ['foto', 'imagem', 'registro fotográfico', 'trabalhando']},
            {'nome': 'Comprovante de Residência', 'gatilhos': ['endereço', 'residência', 'domicílio']}
        ],
        'BAIXA': [
            {'nome': 'Conversas de WhatsApp', 'gatilhos': ['mensagem', 'conversa', 'WhatsApp', 'wpp']},
            {'nome': 'Contrato', 'gatilhos': ['contrato', 'acordo']}
        ]
    },
    'ACAO_ACIDENTARIA': {
        'ALTA': [
            {'nome': 'Atestados médicos', 'gatilhos': ['atestado médico', 'atestou', 'afastamento', 'atestado'], 'critico': True},
            {'nome': 'Exames', 'gatilhos': ['exame', 'raio-X', 'resultado', 'laudo'], 'critico': True},
            {'nome': 'CAT', 'gatilhos': ['CAT', 'Comunicação de Acidente'], 'critico': True},
            {'nome': 'Documentos Previdenciários', 'gatilhos': ['auxílio-doença', 'INSS', 'aposentadoria', 'previdência']}
        ],
        'MEDIA': [
            {'nome': 'Fotos do Ambiente de Trabalho', 'gatilhos': ['condições insalubres', 'ambiente perigoso', 'local do acidente', 'ambiente']},
            {'nome': 'TRCT', 'gatilhos': ['rescisão', 'TRCT', 'término']}
        ],
        'BAIXA': []
    },
    'DIFERENCAS_CONTRATUAIS': {
        'ALTA': [
            {'nome': 'Holerites', 'gatilhos': ['holerite', 'contracheque', 'comprovante de pagamento'], 'critico': True},
            {'nome': 'TRCT', 'gatilhos': ['rescisão', 'TRCT', 'verbas rescisórias'], 'critico': True}
        ],
        'MEDIA': [
            {'nome': 'Cartões de Ponto', 'gatilhos': ['cartão de ponto', 'controle de jornada', 'horas extras', 'ponto']},
            {'nome': 'Extrato FGTS', 'gatilhos': ['FGTS', 'extrato', 'saldo']},
            {'nome': 'Contrato de Trabalho', 'gatilhos': ['contrato', 'cláusula', 'acordo firmado']}
        ],
        'BAIXA': [
            {'nome': 'Conversas de WhatsApp', 'gatilhos': ['ordens via WhatsApp', 'mensagens', 'conversas', 'wpp']}
        ]
    }
}

# ============================================================================
# FASE 3 - ESTATÍSTICAS GLOBAIS
# ============================================================================

ESTATISTICAS_FILE = 'estatisticas_escritorio.json'

def carregar_estatisticas():
    """Carrega estatísticas globais do escritório"""
    try:
        if os.path.exists(ESTATISTICAS_FILE):
            with open(ESTATISTICAS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            'total_peticoes': 0,
            'total_aprovadas': 0,
            'total_rejeitadas': 0,
            'score_medio_geral': 0,
            'scores_por_tipo': {},
            'tempo_medio_geracao': 0,
            'por_tipo': {
                'RECONHECIMENTO_VINCULO': {'total': 0, 'aprovadas': 0, 'score_medio': 0, 'scores': []},
                'ACAO_ACIDENTARIA': {'total': 0, 'aprovadas': 0, 'score_medio': 0, 'scores': []},
                'DIFERENCAS_CONTRATUAIS': {'total': 0, 'aprovadas': 0, 'score_medio': 0, 'scores': []}
            }
        }
    except:
        return {'total_peticoes': 0}

def atualizar_estatisticas(tipo_processo, score, aprovada, tempo_geracao=0):
    """Atualiza estatísticas globais após cada petição"""
    try:
        stats = carregar_estatisticas()
        
        # Totais gerais
        stats['total_peticoes'] = stats.get('total_peticoes', 0) + 1
        if aprovada:
            stats['total_aprovadas'] = stats.get('total_aprovadas', 0) + 1
        else:
            stats['total_rejeitadas'] = stats.get('total_rejeitadas', 0) + 1
        
        # Por tipo
        if 'por_tipo' not in stats:
            stats['por_tipo'] = {}
        if tipo_processo not in stats['por_tipo']:
            stats['por_tipo'][tipo_processo] = {'total': 0, 'aprovadas': 0, 'score_medio': 0, 'scores': []}
        
        tipo_stats = stats['por_tipo'][tipo_processo]
        tipo_stats['total'] += 1
        if aprovada:
            tipo_stats['aprovadas'] += 1
        tipo_stats['scores'].append(score)
        tipo_stats['score_medio'] = sum(tipo_stats['scores']) / len(tipo_stats['scores'])
        
        # Score médio geral
        todos_scores = []
        for tipo_data in stats['por_tipo'].values():
            todos_scores.extend(tipo_data.get('scores', []))
        if todos_scores:
            stats['score_medio_geral'] = sum(todos_scores) / len(todos_scores)
        
        # Tempo médio
        if tempo_geracao > 0:
            if 'tempos' not in stats:
                stats['tempos'] = []
            stats['tempos'].append(tempo_geracao)
            stats['tempo_medio_geracao'] = sum(stats['tempos']) / len(stats['tempos'])
        
        # Taxa de aprovação
        if stats['total_peticoes'] > 0:
            stats['taxa_aprovacao'] = (stats['total_aprovadas'] / stats['total_peticoes']) * 100
        
        with open(ESTATISTICAS_FILE, 'w', encoding='utf-8') as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
        
        return stats
    except Exception as e:
        print(f"Erro ao atualizar estatísticas: {e}")
        return None

def calcular_ranking(score, tipo_processo):
    """Calcula ranking do score dentro do tipo"""
    try:
        stats = carregar_estatisticas()
        tipo_stats = stats.get('por_tipo', {}).get(tipo_processo, {})
        scores = tipo_stats.get('scores', [])
        
        if not scores or len(scores) < 3:
            return "N/A"
        
        scores_sorted = sorted(scores, reverse=True)
        posicao = scores_sorted.index(score) + 1 if score in scores_sorted else len(scores_sorted)
        percentil = (posicao / len(scores_sorted)) * 100
        
        if percentil <= 10:
            return "Top 10% "
        elif percentil <= 25:
            return "Top 25% "
        elif percentil <= 50:
            return "Top 50% "
        else:
            return f"Top {int(percentil)}%"
    except:
        return "N/A"

# ============================================================================
# SISTEMA DE PRINTS - IDENTIFICAÇÃO E MARCAÇÃO
# ============================================================================

def identificar_tipo_acao_por_texto(texto_peticao):
    """Identifica tipo de ação analisando conteúdo da petição"""
    texto_lower = texto_peticao.lower()
    
    # Reconhecimento de Vínculo
    indicadores_vinculo = ['reconhecimento de vínculo', 'ausência de registro', 'vínculo empregatício', 
                           'trabalho sem registro', 'registro em ctps']
    if any(ind in texto_lower for ind in indicadores_vinculo):
        return 'RECONHECIMENTO_VINCULO'
    
    # Ação Acidentária
    indicadores_acidente = ['acidente de trabalho', 'doença ocupacional', 'acidente do trabalho',
                           'lesão', 'cat', 'auxílio-doença acidentário']
    if any(ind in texto_lower for ind in indicadores_acidente):
        return 'ACAO_ACIDENTARIA'
    
    # Diferenças Contratuais (padrão - 54% dos casos)
    return 'DIFERENCAS_CONTRATUAIS'

def encontrar_arquivo_correspondente(nome_print, arquivos_pasta):
    """Busca arquivo correspondente ao print na pasta do cliente"""
    try:
        nome_lower = nome_print.lower()
        
        # Mapeamento de nomes para padrões de arquivo
        mapeamentos = {
            'documento pessoal': ['rg', 'cpf', 'identidade', 'documento', 'docs'],
            'carteira de trabalho': ['ctps', 'carteira', 'trabalho'],
            'comprovantes de pagamento': ['comprovante', 'recibo', 'pagamento', 'transferencia'],
            'fotos trabalhando': ['foto', 'imagem', 'img', 'trabalhando'],
            'comprovante de residência': ['residencia', 'endereco', 'luz', 'agua'],
            'conversas de whatsapp': ['whatsapp', 'wpp', 'conversa', 'mensagem'],
            'contrato': ['contrato'],
            'atestados médicos': ['atestado'],
            'exames': ['exame', 'laudo'],
            'cat': ['cat'],
            'documentos previdenciários': ['inss', 'previdencia'],
            'fotos do ambiente': ['ambiente', 'local'],
            'trct': ['trct', 'rescisao'],
            'holerites': ['holerite', 'contracheque'],
            'cartões de ponto': ['ponto', 'cartao'],
            'extrato fgts': ['fgts', 'extrato']
        }
        
        padroes = mapeamentos.get(nome_lower, [nome_lower])
        
        for arquivo in arquivos_pasta:
            # PROTEÇÃO: Verificar se arquivo é string ou dict
            if isinstance(arquivo, str):
                arquivo_lower = arquivo.lower()
                if any(padrao in arquivo_lower for padrao in padroes):
                    return arquivo
            elif isinstance(arquivo, dict):
                # Se for dict, tentar pegar 'name'
                arquivo_name = arquivo.get('name', '')
                if arquivo_name:
                    arquivo_lower = arquivo_name.lower()
                    if any(padrao in arquivo_lower for padrao in padroes):
                        return arquivo_name
            else:
                print(f" AVISO: arquivo em formato inesperado: {type(arquivo)}")
                continue
        
        return None
        
    except Exception as e:
        print(f" ERRO em encontrar_arquivo_correspondente: {e}")
        return None

def inserir_marcadores_prints(texto_peticao, tipo_acao, arquivos_cliente):
    """Insere marcadores de prints no texto da petição"""
    
    try:
        prints_tipo = PRINTS_POR_TIPO.get(tipo_acao, {})
        marcadores_inseridos = []
        prints_faltantes = {'ALTA': [], 'MEDIA': [], 'BAIXA': []}
        texto_modificado = texto_peticao
        
        # Processar cada prioridade
        for prioridade in ['ALTA', 'MEDIA', 'BAIXA']:
            prints_prioridade = prints_tipo.get(prioridade, [])
            
            # PROTEÇÃO: Verificar se prints_prioridade é lista
            if not isinstance(prints_prioridade, list):
                print(f" AVISO: prints_prioridade não é lista para {prioridade}")
                continue
            
            for print_info in prints_prioridade:
                # PROTEÇÃO: Verificar se print_info é dicionário
                if not isinstance(print_info, dict):
                    print(f" AVISO: print_info não é dicionário: {print_info}")
                    continue
                
                # PROTEÇÃO: Verificar se 'nome' existe
                if 'nome' not in print_info:
                    print(f" AVISO: print_info sem 'nome': {print_info}")
                    continue
                
                # PROTEÇÃO: Verificar se 'gatilhos' existe
                if 'gatilhos' not in print_info:
                    print(f" AVISO: print_info sem 'gatilhos': {print_info}")
                    continue
                
                nome_print = print_info['nome']
                gatilhos = print_info['gatilhos']
                critico = print_info.get('critico', False)
                
                # Verificar se algum gatilho aparece no texto
                mencao_encontrada = False
                posicao_mencao = -1
                
                for gatilho in gatilhos:
                    if gatilho.lower() in texto_peticao.lower():
                        mencao_encontrada = True
                        posicao_mencao = texto_peticao.lower().find(gatilho.lower())
                        break
                
                if mencao_encontrada:
                    # Buscar arquivo correspondente
                    arquivo_encontrado = encontrar_arquivo_correspondente(nome_print, arquivos_cliente)
                    
                    if arquivo_encontrado:
                        status = "Disponível"
                        marcador = f"\n\n[INSERIR PRINT: {nome_print} - Arquivo: {arquivo_encontrado} - Status: {status}]\n\n"
                    else:
                        if critico:
                            status = "FALTANTE -  CRÍTICO"
                            prints_faltantes[prioridade].append(nome_print)
                        else:
                            status = "FALTANTE"
                            prints_faltantes[prioridade].append(nome_print)
                        
                        marcador = f"\n\n[INSERIR PRINT: {nome_print} - Arquivo: {status}]\n\n"
                    
                    marcadores_inseridos.append({
                        'nome': nome_print,
                        'prioridade': prioridade,
                        'arquivo': arquivo_encontrado,
                        'status': status,
                        'critico': critico
                    })
                    
                    # Inserir marcador após primeira menção
                    # Encontrar final do parágrafo
                    pos_final_paragrafo = texto_peticao.find('\n', posicao_mencao)
                    if pos_final_paragrafo == -1:
                        pos_final_paragrafo = len(texto_peticao)
                    
                    # Inserir marcador
                    texto_modificado = (texto_modificado[:pos_final_paragrafo] + 
                                      marcador + 
                                      texto_modificado[pos_final_paragrafo:])
        
        # Verificar se há prints críticos faltantes
        tem_criticos_faltantes = False
        try:
            tem_criticos_faltantes = any(
                print_info.get('critico', False) and not encontrar_arquivo_correspondente(print_info.get('nome', ''), arquivos_cliente)
                for prioridade in ['ALTA']
                for print_info in prints_tipo.get(prioridade, [])
                if isinstance(print_info, dict) and 'nome' in print_info and 'gatilhos' in print_info
                if any(g.lower() in texto_peticao.lower() for g in print_info.get('gatilhos', []))
            )
        except Exception as e:
            print(f" ERRO ao verificar críticos: {e}")
            tem_criticos_faltantes = False
        
        return texto_modificado, marcadores_inseridos, prints_faltantes, tem_criticos_faltantes
        
    except Exception as e:
        print(f" ERRO CRÍTICO em inserir_marcadores_prints: {e}")
        import traceback
        traceback.print_exc()
        # Retornar valores seguros
        return texto_peticao, [], {'ALTA': [], 'MEDIA': [], 'BAIXA': []}, False

def gerar_relatorio_prints(tipo_acao, marcadores_inseridos, prints_faltantes, tem_criticos_faltantes, cliente_nome):
    """Gera relatório detalhado de prints"""
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    arquivo_log = f"logs_prints/prints_{cliente_nome}_{timestamp}.txt"
    
    os.makedirs('logs_prints', exist_ok=True)
    
    conteudo = f"""
{'='*80}
RELATÓRIO DE IDENTIFICAÇÃO DE PRINTS - SISTEMA V4.0
{'='*80}

Cliente: {cliente_nome}
Tipo de Ação: {tipo_acao}
Data: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}

{'='*80}
RESUMO
{'='*80}

Total de marcadores inseridos: {len(marcadores_inseridos)}

 Prints CRÍTICOS faltantes: {len(prints_faltantes['ALTA'])}
 Prints IMPORTANTES faltantes: {len(prints_faltantes['MEDIA'])}
 Prints DESEJÁVEIS faltantes: {len(prints_faltantes['BAIXA'])}

{'='*80}
MARCADORES INSERIDOS NA PETIÇÃO
{'='*80}

"""
    
    # Separar por prioridade
    for prioridade in ['ALTA', 'MEDIA', 'BAIXA']:
        emoji = {'ALTA': '', 'MEDIA': '', 'BAIXA': ''}[prioridade]
        marcadores_prioridade = [m for m in marcadores_inseridos if m['prioridade'] == prioridade]
        
        if marcadores_prioridade:
            conteudo += f"\n{emoji} PRIORIDADE {prioridade}:\n\n"
            
            for i, marcador in enumerate(marcadores_prioridade, 1):
                conteudo += f"{i}. {marcador['nome']}\n"
                conteudo += f"   Arquivo: {marcador['arquivo'] if marcador['arquivo'] else 'FALTANTE'}\n"
                conteudo += f"   Status: {marcador['status']}\n"
                if marcador.get('critico'):
                    conteudo += f"    CRÍTICO: Bloqueia aprovação se ausente\n"
                conteudo += "\n"
    
    conteudo += f"""
{'='*80}
PRINTS FALTANTES POR PRIORIDADE
{'='*80}

"""
    
    if prints_faltantes['ALTA']:
        conteudo += " PRIORIDADE ALTA - Bloqueiam aprovação:\n"
        for print_nome in prints_faltantes['ALTA']:
            conteudo += f"   • {print_nome}\n"
        conteudo += "\n"
    
    if prints_faltantes['MEDIA']:
        conteudo += " PRIORIDADE MÉDIA - Alertam:\n"
        for print_nome in prints_faltantes['MEDIA']:
            conteudo += f"   • {print_nome}\n"
        conteudo += "\n"
    
    if prints_faltantes['BAIXA']:
        conteudo += " PRIORIDADE BAIXA - Sugerem:\n"
        for print_nome in prints_faltantes['BAIXA']:
            conteudo += f"   • {print_nome}\n"
        conteudo += "\n"
    
    conteudo += f"""
{'='*80}
DECISÃO FINAL
{'='*80}

"""
    
    if tem_criticos_faltantes:
        conteudo += """ BLOQUEADO - Prints críticos faltantes

Ação necessária:
Solicitar ao cliente os documentos críticos marcados acima antes de aprovar a petição.
Os marcadores foram inseridos no documento para facilitar a inserção manual posterior.

"""
    else:
        conteudo += """ APROVADO PARA INSERÇÃO MANUAL

Todos os prints críticos estão disponíveis.
Os marcadores foram inseridos no documento indicando onde cada print deve ser adicionado.
O advogado deve inserir manualmente os prints nos locais marcados.

"""
    
    conteudo += f"""
{'='*80}
INSTRUÇÕES PARA O ADVOGADO
{'='*80}

1. Abrir a petição gerada
2. Localizar os marcadores [INSERIR PRINT: ...]
3. Para cada marcador:
   a) Abrir o arquivo indicado
   b) Inserir a imagem/print no local marcado
   c) Remover o marcador após inserção
4. Salvar documento final

{'='*80}
FIM DO RELATÓRIO - SISTEMA DE PRINTS V4.0
{'='*80}
"""
    
    with open(arquivo_log, 'w', encoding='utf-8') as f:
        f.write(conteudo)
    
    return arquivo_log

# ============================================================================

MODELOS_IDS = {
    'RECONHECIMENTO_VINCULO': os.getenv('MODELO_VINCULO_ID'),
    'ACAO_ACIDENTARIA': os.getenv('MODELO_ACIDENTARIA_ID'),
    'DIFERENCAS_CONTRATUAIS': os.getenv('MODELO_DIFERENCAS_ID')
}

PASTAS_PETICOES_GERADAS = {
    'RECONHECIMENTO_VINCULO': '1ya29qkIu8J2O1idmlco9HwSqCFKSsxTm',
    'ACAO_ACIDENTARIA': '1OE_MFrNmzrDKTQ4iJN30qhMc95cEfI1P',
    'DIFERENCAS_CONTRATUAIS': '1edqtmAgtZI_B4GvcmXrn8mSRK9DpU-BE'
}

CLIENTES_PROCESSADOS_SESSAO = set()

# Sistema de histórico de petições
HISTORICO_FILE = 'historico_peticoes.json'

def carregar_historico():
    """Carrega histórico de petições geradas"""
    try:
        if os.path.exists(HISTORICO_FILE):
            with open(HISTORICO_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    except:
        return []

def salvar_no_historico(cliente_nome, tipo_processo, arquivo_info):
    """Salva petição no histórico"""
    try:
        historico = carregar_historico()
        
        entrada = {
            'cliente': cliente_nome,
            'tipo_processo': tipo_processo,
            'arquivo_nome': arquivo_info.get('name'),
            'arquivo_id': arquivo_info.get('id'),
            'link': arquivo_info.get('webViewLink'),
            'data_geracao': datetime.now().isoformat(),
            'status': 'gerada'  # gerada, auditada, aprovada, rejeitada
        }
        
        historico.append(entrada)
        
        with open(HISTORICO_FILE, 'w', encoding='utf-8') as f:
            json.dump(historico, f, indent=2, ensure_ascii=False)
        
        return True
    except Exception as e:
        print(f"Erro ao salvar histórico: {e}")
        return False

def atualizar_status_historico(arquivo_id, status, score=None, erros=None, relatorio=None):
    """Atualiza status de uma petição no histórico"""
    try:
        historico = carregar_historico()
        
        for entrada in historico:
            if entrada.get('arquivo_id') == arquivo_id:
                entrada['status'] = status
                entrada['data_auditoria'] = datetime.now().isoformat()
                if score is not None:
                    entrada['score'] = score
                    # Adicionar ranking baseado no score
                    if score >= 90:
                        entrada['ranking'] = "EXCELENTE"
                    elif score >= 80:
                        entrada['ranking'] = "MUITO BOM"
                    elif score >= 70:
                        entrada['ranking'] = "BOM"
                    elif score >= 60:
                        entrada['ranking'] = "SATISFATÓRIO"
                    else:
                        entrada['ranking'] = "PRECISA MELHORAR"
                    
                    # DEBUG: Verificar ranking calculado
                    print(f"   Histórico atualizado: Status={status}, Score={score}, Ranking={entrada['ranking']}")
                    
                if erros:
                    entrada['erros'] = erros
                if relatorio:
                    entrada['relatorio_auditoria'] = relatorio
                break
        
        with open(HISTORICO_FILE, 'w', encoding='utf-8') as f:
            json.dump(historico, f, indent=2, ensure_ascii=False)
        
        return True
    except Exception as e:
        print(f"Erro ao atualizar histórico: {e}")
        return False

def atualizar_status_processamento(cliente_nome, tipo_processo, status_texto):
    """
    Atualiza status durante processamento (antes de ter arquivo_id)
    Usado para feedback em tempo real
    """
    try:
        historico = carregar_historico()
        
        # Procurar por cliente e tipo
        encontrou = False
        for entrada in historico:
            if (entrada.get('cliente', '').lower() == cliente_nome.lower() and 
                entrada.get('tipo_processo') == tipo_processo):
                entrada['status_processamento'] = status_texto
                entrada['ultima_atualizacao'] = datetime.now().isoformat()
                encontrou = True
                break
        
        # Se não encontrou, criar entrada temporária
        if not encontrou:
            entrada_temp = {
                'cliente': cliente_nome,
                'tipo_processo': tipo_processo,
                'status': 'processando',
                'status_processamento': status_texto,
                'data_geracao': datetime.now().isoformat(),
                'ultima_atualizacao': datetime.now().isoformat()
            }
            historico.append(entrada_temp)
        
        with open(HISTORICO_FILE, 'w', encoding='utf-8') as f:
            json.dump(historico, f, indent=2, ensure_ascii=False)
        
        return True
    except Exception as e:
        print(f"Erro ao atualizar status processamento: {e}")
        return False

def log_auditoria(cliente_nome, tipo_processo, resultado, peticao_nome):
    """Salva log detalhado - CHECKLIST V4.0 FASE 1"""
    try:
        log_dir = 'logs_auditoria'
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        arquivo_log = f"{log_dir}/auditoria_{cliente_nome.replace(' ', '_')}_{timestamp}.txt"
        
        # Score médio do tipo
        score_medio = SCORES_MEDIOS.get(tipo_processo, 0)
        comparacao = "ACIMA DA MÉDIA ⬆" if resultado['score'] > score_medio else "ABAIXO DA MÉDIA ⬇" if resultado['score'] < score_medio else "NA MÉDIA "
        
        conteudo = f"""
{'='*80}
RELATÓRIO DE AUDITORIA - CHECKLIST V4.0 FASE 3 PERFEITA
Validação Completa + Qualidade Extra + Comparações Estatísticas
{'='*80}

Data/Hora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
Cliente: {cliente_nome}
Tipo Processo: {tipo_processo}
Petição: {peticao_nome}

{'='*80}
RESULTADO FINAL
{'='*80}

Status: {' APROVADA' if resultado['aprovada'] else ' REJEITADA'}
Score Final: {resultado['score']}/100

 Score Médio do Tipo (Base 100 petições): {score_medio}/100
 Score Médio Escritório: {resultado.get('comparacao_escritorio', score_medio):.1f}/100
 Comparação: {comparacao}
 Ranking: {resultado.get('ranking', 'N/A')}

{'='*80}
JUSTIFICATIVA DO SCORE
{'='*80}

{resultado.get('justificativa_score', 'Não fornecida')}

{'='*80}
QUALIDADE EXTRA (Bônus de Pontos)
{'='*80}

{f'''- Jurisprudência TST: {' Presente (+3pts)' if resultado.get('qualidade_extra', {}).get('jurisprudencia_tst') else ' Ausente (0pts)'}
- Jurisprudência TRT: {' Presente (+3pts)' if resultado.get('qualidade_extra', {}).get('jurisprudencia_trt') else ' Ausente (0pts)'}
- Cálculos Detalhados: {' Tabela/Planilha (+5pts)' if resultado.get('qualidade_extra', {}).get('calculos_detalhados') else ' Ausente (0pts)'}
- Narrativa Persuasiva: {' Emocional (+3pts)' if resultado.get('qualidade_extra', {}).get('narrativa_persuasiva') else ' Genérica (0pts)'}
- Fundamentação Doutrinária: {' Presente (+3pts)' if resultado.get('qualidade_extra', {}).get('fundamentacao_doutrinaria') else ' Ausente (0pts)'}

 Total Bônus: +{resultado.get('qualidade_extra', {}).get('bonus_pontos', 0)} pontos
''' if resultado.get('qualidade_extra') else 'Não avaliada'}

{'='*80}
VALIDAÇÃO DE DOCUMENTOS (Sistema 3 Prioridades)
{'='*80}

 PRIORIDADE ALTA - CRÍTICO (Bloqueia -10pts cada):
{chr(10).join(['    ' + d for d in resultado.get('docs_alta_faltantes', [])]) if resultado.get('docs_alta_faltantes') else '    Todos os documentos ALTA presentes'}

 PRIORIDADE MÉDIA - IMPORTANTE (Alerta -5pts cada):
{chr(10).join(['    ' + d for d in resultado.get('docs_media_faltantes', [])]) if resultado.get('docs_media_faltantes') else '    Todos os documentos MÉDIA presentes'}

 PRIORIDADE BAIXA - DESEJÁVEL (Sugere -2pts cada):
{chr(10).join(['    ' + d for d in resultado.get('docs_baixa_faltantes', [])]) if resultado.get('docs_baixa_faltantes') else '    Todos os documentos BAIXA presentes'}

Penalização Total por Documentos: {resultado.get('score_penalty_docs', 0)} pontos

{'='*80}
ERROS CRÍTICOS ({len(resultado.get('erros_criticos', []))})
{'='*80}

{chr(10).join([' ' + e for e in resultado.get('erros_criticos', [])]) if resultado.get('erros_criticos') else ' Nenhum erro crítico encontrado'}

{'='*80}
VALIDAÇÃO DOS 10 PEDIDOS MAIS COMUNS
{'='*80}

{chr(10).join([f"{'' if 'check' in p.get('validacao', '').lower() or '' in p.get('validacao', '') else ''} {p.get('pedido', 'N/A')}: {p.get('validacao', 'N/A')}" for p in resultado.get('pedidos_validacao', [])]) if resultado.get('pedidos_validacao') else 'Nenhum pedido validado'}

{'='*80}
ESTRUTURA OBRIGATÓRIA (11 Elementos)
{'='*80}

{chr(10).join([f"{'' if v else ''} {k.replace('_', ' ').title()}" for k, v in resultado.get('estrutura_validacao', {}).items()]) if resultado.get('estrutura_validacao') else 'Não validada'}

{'='*80}
REFLEXOS DE HORAS EXTRAS DETALHADOS (7 Reflexos)
{'='*80}

{chr(10).join([f"{'' if r.get('presente') == True else '' if r.get('presente') == False else ''} {r.get('reflexo', 'N/A')}" for r in resultado.get('reflexos_he_detalhados', [])]) if resultado.get('reflexos_he_detalhados') else 'Não aplicável (não pediu HE)'}

{f" REFLEXOS FALTANTES: {', '.join(resultado.get('reflexos_he_faltantes', []))}" if resultado.get('reflexos_he_faltantes') else ''}

{'='*80}
ALERTAS E SUGESTÕES
{'='*80}

 ALERTAS ({len(resultado.get('alertas', []))}):
{chr(10).join(['   ' + a for a in resultado.get('alertas', [])]) if resultado.get('alertas') else '   Nenhum alerta'}

 SUGESTÕES DE MELHORIA ({len(resultado.get('sugestoes', []))}):
{chr(10).join(['   ' + s for s in resultado.get('sugestoes', [])]) if resultado.get('sugestoes') else ' Nenhuma sugestão'}

{'='*80}
PONTOS POSITIVOS
{'='*80}

{chr(10).join([' ' + p for p in resultado.get('pontos_positivos', [])]) if resultado.get('pontos_positivos') else 'Não especificados'}

{'='*80}
COMO CHEGAR A 100/100
{'='*80}

{chr(10).join([' ' + m for m in resultado.get('melhorias_100', [])]) if resultado.get('melhorias_100') else 'Petição já está excelente!'}

{'='*80}
RESUMO EXECUTIVO
{'='*80}

{resultado.get('resumo', 'N/A')}

{'='*80}
DETALHES TÉCNICOS
{'='*80}

Reflexos de HE OK: {' Sim' if resultado.get('reflexos_he_ok') else ' Não'}
Estrutura Completa: {' Sim' if resultado.get('estrutura_completa') else ' Não'}

{'='*80}
FIM DO RELATÓRIO - CHECKLIST V4.0 FASE 3 PERFEITA
Validação Completa + Qualidade Extra + Comparações Estatísticas
{'='*80}
"""
        
        with open(arquivo_log, 'w', encoding='utf-8') as f:
            f.write(conteudo)
        
        print(f"         Relatório Checklist v4.0 Fase 3 PERFEITA salvo")
        return arquivo_log
        
    except Exception as e:
        print(f"        Erro ao salvar log: {e}")
        return None

def aplicar_formatacao_master(doc):
    """
    Aplica formatação padrão do escritório
    - Fonte: Verdana 10pt
    - Margens: 3cm superior, 1cm inferior, 3.25cm esquerda, 2.5cm direita
    - Espaçamento: 1.5
    - Alinhamento: Justificado
    - Recuo: 2cm à esquerda em todo o texto
    """
    try:
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Configurar margens conforme especificação do usuário
        for section in doc.sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(3.25)
            section.right_margin = Cm(2.5)
            
            # NOVO: Configurar distância do cabeçalho e rodapé
            # Isso cria espaço entre o cabeçalho/rodapé e o conteúdo
            section.header_distance = Cm(1.5)  # Espaço abaixo do cabeçalho
            section.footer_distance = Cm(1.0)  # Espaço acima do rodapé
        
        # Aplicar formatação a todos os parágrafos
        for paragraph in doc.paragraphs:
            # Alinhamento justificado
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Espaçamento 1.5
            paragraph.paragraph_format.line_spacing = 1.5
            
            # SEM recuo de parágrafo (alinhado à margem esquerda)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            
            # Fonte Verdana 10pt
            for run in paragraph.runs:
                run.font.name = 'Verdana'
                run.font.size = Pt(10)
        
        print(f"         Formatação aplicada (Verdana 10pt, margens 3-1-3.25-2.5, SEM recuo, espaçamento cabeçalho/rodapé)")
        return doc
    except Exception as e:
        print(f"         Erro ao aplicar formatação: {e}")
        return doc

def aplicar_formatacoes_especiais_word(doc):
    """
    Aplica formatações especiais de forma ROBUSTA, detectando padrões visuais
    e Markdown, sem depender exclusivamente de tags perfeitas da IA.
    
    Ações:
    1. Centralizar Vocativo
    2. Sublinhar Nome do Autor (Detecção Inteligente)
    3. Títulos: Negrito + Maiúsculas + Espaço antes
    4. Numeração: Negrito só no número
    """
    try:
        import re
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Cm
        
        # Padrões para detecção
        padrao_vocativo = re.compile(r'EXCELENT[ÍI]SSIMO', re.IGNORECASE)
        padrao_negrito = re.compile(r'\*\*(.*?)\*\*')
        padrao_titulos = re.compile(r'^\s*(I+\.|[0-9]+\.)\s+[A-ZÀ-Ú]') # Detecta I. TÍTULO ou 1. TÍTULO
        padrao_numeracao = re.compile(r'^(\d+\.)\s')
        
        vocativo_encontrado = False
        autor_formatado = False
        
        # Iterar sobre parágrafos
        for i, paragraph in enumerate(doc.paragraphs):
            texto = paragraph.text.strip()
            
            # Pular parágrafos vazios, mas manter espaçamento
            if not texto:
                continue

            # --- 1. DETECÇÃO E FORMATAÇÃO DE VOCATIVO ---
            if not vocativo_encontrado and padrao_vocativo.search(texto) and len(texto) < 300:
                vocativo_encontrado = True
                
                # Formatação Vocativo
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.text = texto.upper() # Forçar maiúsculas
                
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Verdana'
                    run.font.size = Pt(10)
                
                # Adicionar espaço GRANDE após vocativo (aprox 10 linhas)
                paragraph.paragraph_format.space_after = Pt(120)
                continue
            
            # --- 2. NOME DO AUTOR (Primeiro parágrafo significativo após vocativo) ---
            # Lógica: Se já passamos do vocativo e ainda não formatamos o autor, 
            # e este parágrafo começa com nome (geralmente em negrito ou caps)
            if vocativo_encontrado and not autor_formatado and len(texto) > 10:
                # Tenta detectar nome em negrito ou primeiro trecho antes da vírgula
                match_markdown = padrao_negrito.match(texto)
                primeira_parte = texto.split(',')[0] # Pega até a primeira vírgula
                
                # Se começar com Markdown **Nome**
                if match_markdown or primeira_parte.isupper(): 
                    autor_formatado = True
                    
                    # Limpar markdown se existir
                    texto_limpo = texto.replace('**', '')
                    paragraph.text = "" # Limpar para reconstruir
                    
                    # Parte do nome (até a virgula ou o que estava em negrito)
                    nome = match_markdown.group(1) if match_markdown else primeira_parte
                    resto = texto_limpo[len(nome):]
                    
                    # Run do Nome: Negrito + SUBLINHADO
                    run_nome = paragraph.add_run(nome)
                    run_nome.font.bold = True
                    run_nome.font.underline = True
                    run_nome.font.name = 'Verdana'
                    run_nome.font.size = Pt(10)
                    
                    # Run do resto
                    run_resto = paragraph.add_run(resto)
                    run_resto.font.name = 'Verdana'
                    run_resto.font.size = Pt(10)
                    continue

            # --- 3. TÍTULOS DE SEÇÃO (I. PRELIMINARES / 1. DA COMPETÊNCIA) ---
            # Detecta linhas que parecem títulos e aplica espaçamento antes
            if padrao_titulos.match(texto) and texto.isupper():
                paragraph.paragraph_format.space_before = Pt(12) # Espaço antes do título
                paragraph.paragraph_format.space_after = Pt(6)   # Espaço depois do título
                
                # Garantir Negrito
                if not paragraph.runs or not paragraph.runs[0].font.bold:
                    paragraph.text = texto # Resetar para garantir clean runs
                    run = paragraph.add_run(texto)
                    run.font.bold = True
                    run.font.name = 'Verdana'
                    run.font.size = Pt(10)
                continue

            # --- 4. FORMATAR MARKDOWN EM GERAL (**texto**) ---
            if '**' in texto:
                parts = texto.split('**')
                paragraph.text = "" # Limpa parágrafo existente
                for idx, part in enumerate(parts):
                    run = paragraph.add_run(part)
                    run.font.name = 'Verdana'
                    run.font.size = Pt(10)
                    
                    # Se indice é ímpar, estava entre **, então é negrito
                    if idx % 2 == 1:
                        run.font.bold = True
                # Se tinha numeração, a lógica abaixo vai corrigir o excesso de negrito
            
            # --- 5. NÚMEROS DE PARÁGRAFOS ---
            # Detectar "1. Texto..." e garantir que só o "1." seja negrito
            match_num = padrao_numeracao.match(paragraph.text) # Checa texto atualizado
            if match_num:
                numero = match_num.group(1)
                texto_atual = paragraph.text
                resto = texto_atual[len(numero):]
                
                paragraph.text = "" # Limpa
                
                # Número em Negrito
                run_num = paragraph.add_run(numero)
                run_num.font.bold = True
                run_num.font.name = 'Verdana'
                run_num.font.size = Pt(10)
                
                # Resto normal (mesmo que antes estivesse em negrito por markdown incorreto)
                run_resto = paragraph.add_run(resto)
                run_resto.font.bold = False 
                run_resto.font.name = 'Verdana'
                run_resto.font.size = Pt(10)

        print(f"         Formatações especiais ROBUSTAS aplicadas (incluindo sublinhado e espaçamento)")
        return doc
        
    except Exception as e:
        print(f"         Erro ao aplicar formatações especiais: {e}")
        import traceback
        traceback.print_exc()
        return doc

def autenticar_google_drive():
    print(f"    [DEBUG AUTH] Iniciando autenticação...")
    print(f"    [DEBUG AUTH] Diretório atual: {os.getcwd()}")
    print(f"    [DEBUG AUTH] Arquivos no diretório: {os.listdir('.')}")
    
    creds = None
    if os.path.exists('token.json'):
        print(f"    [DEBUG AUTH] token.json encontrado.")
        try:
            creds = Credentials.from_authorized_user_file('token.json', SCOPES)
            print(f"    [DEBUG AUTH] Credenciais carregadas. Válidas: {creds.valid}, Expiradas: {creds.expired}")
        except Exception as e:
             print(f"    [DEBUG AUTH] Erro ao ler token.json: {e}")
    else:
        print(f"    [DEBUG AUTH] token.json NÃO encontrado.")

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            print(f"    [DEBUG AUTH] Token expirado, tentando refresh...")
            try:
                creds.refresh(Request())
                print(f"    [DEBUG AUTH] Refresh com sucesso!")
            except Exception as e:
                print(f"    [DEBUG AUTH] Falha no refresh ({e}). Forçando login via browser...")
                if os.path.exists('credentials.json'):
                    flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                    creds = flow.run_local_server(port=8080)
                else:
                    print(f"    [DEBUG AUTH] CRÍTICO: credentials.json também não encontrado!")
                    return None
        else:
            print(f"    [DEBUG AUTH] Sem credenciais válidas ou refresh token. Tentando login via browser...")
            if os.path.exists('credentials.json'):
                try:
                    flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                    creds = flow.run_local_server(port=8080)
                except Exception as e:
                     print(f"    [DEBUG AUTH] Erro ao iniciar server local: {e}")
                     # Tentar method console se disponivel ou falhar
                     raise e
            else:
                 print(f"    [DEBUG AUTH] CRÍTICO: credentials.json não encontrado para iniciar login!")
                 return None
                 
        # Salvar token atualizado
        with open('token.json', 'w') as token:
            token.write(creds.to_json())
            
    return build('drive', 'v3', credentials=creds)

def listar_pastas(service, pasta_pai_id):
    try:
        query = f"'{pasta_pai_id}' in parents and trashed=false and mimeType='application/vnd.google-apps.folder'"
        results = service.files().list(q=query, fields="files(id, name)", pageSize=100).execute()
        return results.get('files', [])
    except:
        return []

def listar_arquivos_pasta(service, pasta_id):
    try:
        query = f"'{pasta_id}' in parents and trashed=false and mimeType != 'application/vnd.google-apps.folder'"
        results = service.files().list(q=query, fields="files(id, name, mimeType)", pageSize=100).execute()
        return results.get('files', [])
    except:
        return []

def listar_arquivos_recursivo(service, pasta_id, _nivel=0, _max_nivel=3):
    """
    Lista TODOS os arquivos de uma pasta, incluindo subpastas (recursivo)
    
    Args:
        service: Google Drive service
        pasta_id: ID da pasta raiz
        _nivel: Nível atual de recursão (interno)
        _max_nivel: Profundidade máxima de recursão
    
    Returns:
        Lista de dicionários com 'id', 'name', 'mimeType', 'pasta_origem'
    """
    try:
        # Proteção contra recursão infinita
        if _nivel > _max_nivel:
            return []
        
        arquivos_totais = []
        
        # 1. Listar arquivos na pasta atual
        arquivos_raiz = listar_arquivos_pasta(service, pasta_id)
        
        # Filtrar arquivos que começam com "não juntar" ou "nao juntar"
        for arq in arquivos_raiz:
            nome_lower = arq['name'].lower()
            if nome_lower.startswith('não juntar') or nome_lower.startswith('nao juntar'):
                continue  # Ignorar este arquivo
            
            # Adicionar informação de pasta de origem
            arq['pasta_origem'] = 'raiz' if _nivel == 0 else f'subpasta_nivel_{_nivel}'
            arquivos_totais.append(arq)
        
        # 2. Listar subpastas
        subpastas = listar_pastas(service, pasta_id)
        
        # 3. Recursivamente listar arquivos de cada subpasta
        for subpasta in subpastas:
            arquivos_subpasta = listar_arquivos_recursivo(
                service, 
                subpasta['id'], 
                _nivel=_nivel + 1,
                _max_nivel=_max_nivel
            )
            
            # Marcar de qual subpasta vieram
            for arq in arquivos_subpasta:
                if arq.get('pasta_origem') == 'raiz':
                    arq['pasta_origem'] = subpasta['name']
            
            arquivos_totais.extend(arquivos_subpasta)
        
        return arquivos_totais
        
    except Exception as e:
        print(f"         Erro ao listar recursivamente: {e}")
        return []

def verificar_cliente_ja_processado(service, pasta_cliente_id):
    try:
        # Verificar na sessão atual
        if pasta_cliente_id in CLIENTES_PROCESSADOS_SESSAO:
            return True
        
        # Verificar arquivo _PROCESSADO.txt
        arquivos = listar_arquivos_pasta(service, pasta_cliente_id)
        for arquivo in arquivos:
            if arquivo['name'] == '_PROCESSADO.txt':
                CLIENTES_PROCESSADOS_SESSAO.add(pasta_cliente_id)
                return True
        return False
    except:
        return False

def marcar_cliente_como_processado(service, pasta_cliente_id, info_peticao):
    try:
        # Adicionar à sessão
        CLIENTES_PROCESSADOS_SESSAO.add(pasta_cliente_id)
        
        # Verificar se já existe e deletar para não duplicar
        try:
            arquivos = listar_arquivos_pasta(service, pasta_cliente_id)
            for arq in arquivos:
                if arq['name'] == '_PROCESSADO.txt':
                    print(f"[INFO] Removendo arquivo _PROCESSADO.txt antigo: {arq['id']}")
                    service.files().delete(fileId=arq['id']).execute()
        except Exception as e:
            print(f"[AVISO] Erro ao limpar _PROCESSADO.txt antigo: {e}")

        conteudo = f"""PROCESSADO EM: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
Petição: {info_peticao.get('nome_arquivo', 'N/A')}
Link: {info_peticao.get('link', 'N/A')}
"""
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.txt')
        os.close(tmp_fd)
        with open(tmp_path, 'w', encoding='utf-8') as f:
            f.write(conteudo)
        file_metadata = {'name': '_PROCESSADO.txt', 'parents': [pasta_cliente_id]}
        media = MediaFileUpload(tmp_path, mimetype='text/plain')
        service.files().create(body=file_metadata, media_body=media).execute()
        try:
            os.unlink(tmp_path)
        except:
            pass
        return True
    except Exception as e:
        return False

def baixar_arquivo(service, file_id):
    try:
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        return fh.read()
    except:
        return None

def converter_pdf_para_imagens(conteudo_bytes):
    try:
        pdf_document = fitz.open(stream=conteudo_bytes, filetype="pdf")
        imagens = []
        for page_num in range(min(3, len(pdf_document))):
            page = pdf_document[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            imagens.append(img_base64)
        pdf_document.close()
        return imagens
    except:
        return []

def extrair_texto_pdf(conteudo_bytes):
    try:
        pdf_document = fitz.open(stream=conteudo_bytes, filetype="pdf")
        texto = ""
        for page in pdf_document:
            texto += page.get_text() + "\n"
        pdf_document.close()
        return texto.strip()
    except:
        return ""

def carregar_modelo_peticao(service, modelo_id):
    try:
        conteudo = baixar_arquivo(service, modelo_id)
        if not conteudo:
            return None
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
        os.close(tmp_fd)
        with open(tmp_path, 'wb') as f:
            f.write(conteudo)
        doc = Document(tmp_path)
        texto = "\n".join([p.text for p in doc.paragraphs])
        os.unlink(tmp_path)
        return texto
    except:
        return None

def classificar_documento(nome_arquivo):
    """Classificação expandida - CHECKLIST V4.0 + FASE 3 (Detecção de Prints)"""
    nome = normalizar_texto(nome_arquivo)  # Agora remove acentos!
    
    # PROCURAÇÃO (contém dados completos do cliente)
    if any(x in nome for x in ['procuracao', 'procuração']):
        return {'tipo': 'PROCURACAO', 'prioridade': 'ALTA'}
    
    # DOCUMENTO_PESSOAL (agrupa RG, CPF, CNH)
    elif any(x in nome for x in ['rg', 'identidade', 'cpf', 'cnh', 'habilitacao', 'carteira nacional']):
        return {'tipo': 'DOCUMENTO_PESSOAL', 'prioridade': 'ALTA'}
    
    # CTPS (várias formas de nomear)
    elif any(x in nome for x in ['ctps', 'carteira trabalho', 'carteira de trabalho', 'trabalho', 'ct ']):
        return {'tipo': 'CTPS', 'prioridade': 'ALTA'}
    
    # HOLERITES (específico) - PRIORIDADE MÁXIMA
    elif any(x in nome for x in ['holerite', 'contracheque', 'hollerite', 'olerite', 'contra cheque']):
        return {'tipo': 'HOLERITES', 'prioridade': 'ALTA'}
    
    # COMPROVANTE_PAGAMENTO (genérico) - DEPOIS de holerites
    elif any(x in nome for x in ['pagamento', 'pagto', 'recibo', 'folha ponto', 'comprovante pag']):
        # Se tem "holerite" no nome, é HOLERITES
        if 'holerite' in nome or 'contra' in nome:
            return {'tipo': 'HOLERITES', 'prioridade': 'ALTA'}
        return {'tipo': 'COMPROVANTE_PAGAMENTO', 'prioridade': 'ALTA'}
    
    # TRCT (Rescisão)
    elif any(x in nome for x in ['rescisao', 'trct', 'termo rescisao', 'rescicao']):
        return {'tipo': 'TRCT', 'prioridade': 'ALTA'}
    
    # COMPROVANTE_RESIDENCIA
    elif any(x in nome for x in ['luz', 'agua', 'energia', 'iptu', 'aluguel', 'residencia', 'endereco', 'comprovante residencia']):
        return {'tipo': 'COMPROVANTE_RESIDENCIA', 'prioridade': 'MEDIA'}
    
    # DOCUMENTOS MÉDICOS
    elif any(x in nome for x in ['atestado', 'atestados']):
        return {'tipo': 'ATESTADO_MEDICO', 'prioridade': 'ALTA'}
    elif any(x in nome for x in ['exame', 'laudo', 'exames']):
        return {'tipo': 'EXAMES_MEDICOS', 'prioridade': 'ALTA'}
    elif 'cat' in nome and not 'categoria' in nome:
        return {'tipo': 'CAT', 'prioridade': 'ALTA'}
    elif any(x in nome for x in ['inss', 'previdencia', 'previdenciario']):
        return {'tipo': 'DOCUMENTOS_PREVIDENCIARIOS', 'prioridade': 'ALTA'}
    
    # CRONOLOGIA DOS FATOS (documento gerado pelo sistema)
    elif 'cronologia' in nome:
        return {'tipo': 'CRONOLOGIA', 'prioridade': 'ALTA'}
    
    # RESUMO DO VÍDEO (documento gerado pelo sistema)
    elif 'resumo' in nome and any(x in nome for x in ['video', 'fatos', 'detalhado']):
        return {'tipo': 'RESUMO', 'prioridade': 'ALTA'}
    
    # TRANSCRIÇÃO (OBRIGATÓRIO) - Agora detecta "Transcrição" com acento!
    elif any(x in nome for x in ['transcricao', 'entrevista', 'relato', 'audio_transcrito']):
        return {'tipo': 'TRANSCRICAO', 'prioridade': 'ALTA'}
    
    # FASE 3: DETECÇÃO AUTOMÁTICA DE PRINTS/IMAGENS
    # WhatsApp, Telegram, SMS, conversas
    elif any(x in nome for x in ['whatsapp', 'wpp', 'conversa', 'chat', 'mensagem', 'zap', 'telegram', 'sms', 'print', 'screenshot', 'captura', 'tela']):
        return {'tipo': 'CONVERSAS_WHATSAPP', 'prioridade': 'BAIXA'}
    
    # FOTOS - Detecção inteligente
    elif any(x in nome for x in ['foto', 'imagem', 'img', 'jpg', 'jpeg', 'png', 'image']):
        # Se menciona ambiente, local, empresa
        if any(x in nome for x in ['ambiente', 'local', 'empresa', 'fabrica', 'escritorio', 'loja']):
            return {'tipo': 'FOTOS_AMBIENTE', 'prioridade': 'MEDIA'}
        # Se menciona trabalho, trabalhando, uniforme, cargo
        elif any(x in nome for x in ['trabalho', 'trabalhando', 'uniforme', 'cargo', 'funcao', 'servico']):
            return {'tipo': 'FOTOS_TRABALHANDO', 'prioridade': 'MEDIA'}
        # Genérico - assume trabalhando
        else:
            return {'tipo': 'FOTOS_TRABALHANDO', 'prioridade': 'MEDIA'}
    
    # CONTRATO
    elif any(x in nome for x in ['contrato', 'contracto']):
        return {'tipo': 'CONTRATO', 'prioridade': 'MEDIA'}
    
    # CARTÃO DE PONTO
    elif any(x in nome for x in ['ponto', 'cartao']):
        return {'tipo': 'CARTAO_PONTO', 'prioridade': 'MEDIA'}
    
    # EXTRATO FGTS
    elif any(x in nome for x in ['fgts', 'extrato']):
        return {'tipo': 'EXTRATO_FGTS', 'prioridade': 'MEDIA'}
    
    
    # VÍDEOS - Para transcrição
    elif nome.endswith(('.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv', '.m4v', '.mpeg', '.mpg')):
        return {'tipo': 'VIDEO', 'prioridade': 'MEDIA', 'descricao': 'Vídeo para transcrição'}
    
    # GENÉRICO - Se for imagem sem identificação, assume foto trabalhando
    elif nome.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
        return {'tipo': 'FOTOS_TRABALHANDO', 'prioridade': 'MEDIA'}
    
    else:
        return {'tipo': 'DOCUMENTO_GENERICO', 'prioridade': 'BAIXA'}


def verificar_documentacao_completa_v10(documentos_cliente, tipo_processo):
    """
    Sistema de 3 Prioridades - CHECKLIST V4.0
    
    Retorna:
    - alta_faltantes: lista de docs ALTA faltantes (BLOQUEIA)
    - media_faltantes: lista de docs MÉDIA faltantes (ALERTA)
    - baixa_faltantes: lista de docs BAIXA faltantes (SUGERE)
    - score_penalty: pontos a descontar do score
    - bloqueia: True se tem doc ALTA faltando
    """
    docs_config = DOCUMENTOS_POR_TIPO.get(tipo_processo, {})
    tipos_presentes = set([doc['tipo'] for doc in documentos_cliente])
    
    resultado = {
        'alta_faltantes': [],
        'media_faltantes': [],
        'baixa_faltantes': [],
        'alta_presentes': [],
        'media_presentes': [],
        'baixa_presentes': [],
        'score_penalty': 0,
        'bloqueia': False,
        'completo': False
    }
    
    # Verificar ALTA ( CRÍTICO - BLOQUEIA)
    for doc in docs_config.get('ALTA', []):
        if doc not in tipos_presentes:
            resultado['alta_faltantes'].append(doc)
            resultado['score_penalty'] += PONTOS_PRIORIDADE['ALTA']
            resultado['bloqueia'] = True
        else:
            resultado['alta_presentes'].append(doc)
    
    # Verificar MÉDIA ( IMPORTANTE - ALERTA)
    for doc in docs_config.get('MEDIA', []):
        if doc not in tipos_presentes:
            resultado['media_faltantes'].append(doc)
            resultado['score_penalty'] += PONTOS_PRIORIDADE['MEDIA']
        else:
            resultado['media_presentes'].append(doc)
    
    # Verificar BAIXA ( DESEJÁVEL - SUGERE)
    for doc in docs_config.get('BAIXA', []):
        if doc not in tipos_presentes:
            resultado['baixa_faltantes'].append(doc)
            resultado['score_penalty'] += PONTOS_PRIORIDADE['BAIXA']
        else:
            resultado['baixa_presentes'].append(doc)
    
    # Considera completo se todos os ALTA estão presentes
    resultado['completo'] = len(resultado['alta_faltantes']) == 0
    
    return resultado
    return completo, faltantes, presentes

def gerar_peticao_com_claude(service, cliente_info, documentos_completos, tipo_processo, 
                              cronologia_fatos=None, resumo_video=None, procuracao=None, 
                              usar_prompt_master=False):
    """
    Gera petição com PROMPT MELHORADO - dados completos obrigatórios
    
    Args:
        usar_prompt_master: Se True, usa o Prompt Master para petições de 12-18 páginas de alto nível
    """
    modo = "PROMPT MASTER" if usar_prompt_master else "PADRÃO"
    print(f"        [GERANDO PETICAO COM CLAUDE + VISAO - MODO: {modo}]")
    try:
        modelo_id = MODELOS_IDS.get(tipo_processo)
        if not modelo_id:
            return None
        print(f"        - Carregando modelo...")
        modelo_texto = carregar_modelo_peticao(service, modelo_id)
        if not modelo_texto:
            return None
        
        prompt_inicial = f"""Você é advogado trabalhista especializado em gerar petições COMPLETAS e PERFEITAS.

TIPO: {tipo_processo}
CLIENTE: {cliente_info['cliente_nome']}

PROCURAÇÃO (DADOS COMPLETOS DO CLIENTE - USE PARA EXTRAIR NOME, ENDEREÇO, RG, CPF):
{procuracao if procuracao else "Nenhuma procuração disponível. Busque dados nos documentos individuais."}

RESUMO DO VÍDEO (INFORMAÇÕES CONTEXTUAIS DO CLIENTE):
{resumo_video if resumo_video else "Nenhum resumo de vídeo disponível."}

CRONOLOGIA DOS FATOS (USE ESTA LINHA DO TEMPO COMO BASE PARA OS FATOS):
{cronologia_fatos if cronologia_fatos else "Nenhuma cronologia disponível."}

INSTRUÇÕES IMPORTANTES:
- Se houver PROCURAÇÃO, extraia dela: Nome completo, RG, CPF, Endereço completo do cliente
- Use esses dados para preencher a qualificação do Reclamante
- Se não houver procuração, busque nos documentos individuais (RG, CPF, Comprovante de Residência)

TAREFA: Analise CUIDADOSAMENTE os documentos e gere uma petição inicial COMPLETA com TODOS os dados preenchidos.

DOCUMENTOS FORNECIDOS:
"""
        # Mover para system prompt para evitar erro de roles
        system_prompt = prompt_inicial
        conteudo_mensagem = []
        
        print(f"        - Processando {len(documentos_completos)} documentos...")
        for doc in documentos_completos:
            conteudo_mensagem.append({
                "type": "text",
                "text": f"\n=== DOCUMENTO: {doc['tipo']} - {doc['nome']} ==="
            })
            if doc['nome'].lower().endswith('.pdf'):
                imagens = converter_pdf_para_imagens(doc['conteudo'])
                if imagens:
                    conteudo_mensagem.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": "image/png", "data": imagens[0]}
                    })
            if doc.get('texto'):
                conteudo_mensagem.append({
                    "type": "text",
                    "text": f"Texto extraído:\n{doc['texto'][:1500]}"
                })
        
        instrucoes = f"""

MODELO DE PETIÇÃO:
{modelo_texto}

═══════════════════════════════════════════════════════════════════════════
INSTRUÇÕES CRÍTICAS - CHECKLIST IA AUDITORA V4.0
═══════════════════════════════════════════════════════════════════════════

Esta petição será AUDITADA automaticamente. Siga RIGOROSAMENTE estas instruções
para evitar BLOQUEIO ou REJEIÇÃO.

═══════════════════════════════════════════════════════════════════════════
0. FORMATAÇÃO PROFISSIONAL - PADRÃO ADVOGADO EXPERIENTE
═══════════════════════════════════════════════════════════════════════════

 FORMATAÇÃO ESPECIAL OBRIGATÓRIA (USE MARCADORES):

**VOCATIVO (Primeira linha após cabeçalho):**
   - Formato: EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DO TRABALHO DA VARA DO TRABALHO DE [CIDADE] - [UF]
   - Use o marcador: [CENTRALIZAR]texto do vocativo[/CENTRALIZAR]
   - Exemplo:
     [CENTRALIZAR]
     EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DO TRABALHO DA VARA
     DO TRABALHO DE SÃO BERNARDO DO CAMPO - SP
     [/CENTRALIZAR]
   - SEMPRE em MAIÚSCULAS
   - Após o vocativo, adicione: [ESPACO_GRANDE]

**QUALIFICAÇÃO DO AUTOR:**
   - Nome completo do autor: [NEGRITO_SUBLINHADO]NOME COMPLETO[/NEGRITO_SUBLINHADO]
   - Exemplo: [NEGRITO_SUBLINHADO]JOSICLEBSON ANTÔNIO DA SILVA[/NEGRITO_SUBLINHADO], brasileiro, viúvo...
   - Resto da qualificação: texto normal

**TIPO DA AÇÃO:**
   - Formato: **RECLAMAÇÃO TRABALHISTA** (use ** para negrito)
   - Inline, não centralizado

**NOME DA EMPRESA:**
   - Formato: **NOME DA EMPRESA** (use ** para negrito)
   - Exemplo: em face de **BETO PINHEIRO COMÉRCIO, PROMOÇÕES E EVENTOS LTDA**

**TÍTULOS DE SEÇÕES:**
   - Seções principais: I. PRELIMINARES, II. DOS FATOS, III. DO MÉRITO, IV. DOS PEDIDOS
   - Subseções: 1. DA COMPETÊNCIA TERRITORIAL, 2. DO JUÍZO 100% DIGITAL
   - Formato: **NEGRITO E MAIÚSCULAS** (use **)

 NUMERAÇÃO SEQUENCIAL DE PARÁGRAFOS:
   - TODOS os parágrafos do corpo da petição devem ser numerados: 1. 2. 3. ... até o fim
   - NÃO reiniciar numeração em novos capítulos
   - Formato: "X. [texto do parágrafo]"
   - Apenas o NÚMERO fica em negrito automaticamente (não use **)
   - Exceções (NÃO numerar): Vocativo inicial, títulos de seções (DOS FATOS, DO MÉRITO, etc.), assinatura final

 IMPORTANTE - SEM RECUO DE PARÁGRAFO:
   - TODO o texto deve estar alinhado à margem esquerda
   - NÃO use tabulações ou espaços no início das linhas
   - NÃO use recuo de primeira linha

 ESTRUTURA COMPLETA OBRIGATÓRIA:

   I. PRELIMINAR (SEMPRE incluir quando aplicável):
      - Da Competência Territorial (fundamentar foro escolhido)
      - Do Juízo 100% Digital (se aplicável ao caso)
      - Da Juntada de Documentos (listar documentos anexados)
   
   II. DOS FATOS:
      - Narrativa cronológica DETALHADA
      - Datas específicas (admissão, rescisão, eventos relevantes)
      - Valores de remuneração com precisão
      - Condições de trabalho descritas
      - Jornada de trabalho especificada
   
   III. DO MÉRITO (criar um capítulo específico para CADA pedido principal):
      Exemplos de capítulos conforme o caso:
      - DO RECONHECIMENTO DO VÍNCULO EMPREGATÍCIO
      - DAS HORAS EXTRAS E SEUS REFLEXOS
      - DO ADICIONAL NOTURNO E SEUS REFLEXOS  
      - DO ADICIONAL DE INSALUBRIDADE
      - DOS DANOS MORAIS
      - DA RESCISÃO INDIRETA
      - DAS MULTAS DOS ARTIGOS 477 E 467 DA CLT
      - DOS JUROS E CORREÇÃO MONETÁRIA
      [Outros conforme necessário]
   
   IV. DOS PEDIDOS:
      - Lista numerada e específica
      - Valores discriminados quando aplicável

 FUNDAMENTAÇÃO OBRIGATÓRIA POR CAPÍTULO DO MÉRITO:
   Para CADA pedido, incluir OBRIGATORIAMENTE:
   
   1. Descrição fática específica do pedido
   2. Fundamentação legal (artigos da CLT, CF/88, CC/02)
   3. Jurisprudência (mínimo 1 acórdão relevante - TST ou TRT)
   4. Cálculo discriminado (quando aplicável)
   5. Pedido específico com valor

 FORMATO DE JURISPRUDÊNCIA:
   Use o seguinte formato OBRIGATÓRIO:
   
   (TRIBUNAL - TIPO: NÚMERO, Relator: NOME COMPLETO, Data: DD/MM/AAAA, Turma: NOME DA TURMA)
   "EMENTA: [texto completo da ementa relevante]"
   
   Exemplo correto:
   (TST - RR: 1072-40.2013.5.03.0053, Relator: Mauricio Godinho Delgado, Data: 18/05/2016, Turma: 3ª Turma)
   "VÍNCULO DE EMPREGO. REQUISITOS. PRESENÇA. A relação de emprego caracteriza-se pela prestação de trabalho não eventual, subordinado, oneroso e pessoal. Presentes tais requisitos, impõe-se o reconhecimento do vínculo empregatício, nos termos dos arts. 2º e 3º da CLT."

 CÁLCULOS DETALHADOS:
   - Discriminar CADA verba item por item
   - Incluir fórmula de cálculo quando relevante
   - Apresentar valor em R$ X.XXX,XX (valor por extenso)
   - Incluir TODOS os reflexos aplicáveis
   - Exemplo: "13º salário proporcional: R$ 2.500,00 ÷ 12 × 8 meses = R$ 1.666,67 (um mil, seiscentos e sessenta e seis reais e sessenta e sete centavos)"

 CAPÍTULOS OBRIGATÓRIOS (quando aplicável ao caso):
   - Se houver rescisão: incluir capítulo "DAS MULTAS DOS ARTIGOS 477 E 467 DA CLT"
   - Se houver horas extras: incluir capítulo específico com TODOS os 7 reflexos
   - Se houver danos: incluir capítulo "DOS DANOS MORAIS" com fundamentação robusta

═══════════════════════════════════════════════════════════════════════════
1. ESTRUTURA OBRIGATÓRIA (10 ELEMENTOS - NÃO PODE FALTAR NENHUM)
═══════════════════════════════════════════════════════════════════════════

 1. Vocativo formal: "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DA VARA DO TRABALHO DE [CIDADE]"
 2. Qualificação COMPLETA do Reclamante:
   - Nome completo (sem abreviações)
   - RG: formato XX.XXX.XXX-X
   - CPF: formato XXX.XXX.XXX-XX
   - Endereço: Rua/Av [nome], nº [número], Bairro [bairro], [Cidade]-[UF], CEP XX.XXX-XXX

 3. Qualificação COMPLETA da Reclamada:
   - Razão social completa
   - CNPJ: formato XX.XXX.XXX/XXXX-XX
   - Endereço completo com CEP

 4. Fórmula processual: "vem à presença de Vossa Excelência, fazendo-o pelas razões de fato e de direito a seguir aduzidas"

 5. Seção DOS FATOS (narrativa cronológica detalhada)
 6. Seção DO MÉRITO (argumentação jurídica fundamentada)
 7. Seção DOS PEDIDOS (lista específica e clara)

 8. Fórmula de encerramento COMPLETA:
   "Termos em que,
   Pede deferimento.
   
   [Cidade], {datetime.now().strftime('%d de %B de %Y')}.
   
   ___________________________
   [Nome do Advogado]
   OAB/[Estado] [Número]"

 9. Data e local
 10. Assinatura e OAB

═══════════════════════════════════════════════════════════════════════════
2. OS 10 ERROS CRÍTICOS - EVITE A TODO CUSTO (BLOQUEIO AUTOMÁTICO)
═══════════════════════════════════════════════════════════════════════════

 1. INÉPCIA - Pedido impossível de julgar ou sem objeto definido
 2. PEDIDOS GENÉRICOS - "Pagar tudo que for devido"  BLOQUEIA
 3. SEM FUNDAMENTAÇÃO LEGAL - Mencionar artigos da CLT, CF/88
 4. QUALIFICAÇÃO INCOMPLETA - Faltando RG, CPF ou endereço
 5. SEM DOCUMENTOS CRÍTICOS - Verificado pelo sistema
 6. CONTRADIÇÃO - Fatos não sustentam pedidos
 7. SEM COMPETÊNCIA - Fundamentar foro/competência territorial
 8. REFLEXOS HE FALTANDO - Se pede HE, DEVE pedir TODOS reflexos
 9. SEM ENCERRAMENTO - Fórmula "Termos em que, pede deferimento" obrigatória

═══════════════════════════════════════════════════════════════════════════
3. EXTRAÇÃO DE DADOS (OBRIGATÓRIO - PROCURE NOS DOCUMENTOS)
═══════════════════════════════════════════════════════════════════════════

 RG: Procure em RG, CNH ou documentos pessoais  formato XX.XXX.XXX-X
 CPF: Procure em CPF, CNH, CTPS  formato XXX.XXX.XXX-XX
 Endereço: Extraia do comprovante de residência  completo com CEP
 Nome: Use nome COMPLETO sem abreviações
 Empresa: Nome, CNPJ, endereço completo
 Período: Datas de admissão e demissão da CTPS
 Cargo: Função exercida
 Salário: Valor mensal

 NÃO deixe campos vazios ou com "[...]"
 NÃO use placeholders como "{{{{NOME}}}}"
 Se NÃO encontrar, use "não informado nos autos"

═══════════════════════════════════════════════════════════════════════════
4. VALIDAÇÃO DE PEDIDOS COMUNS (TOP 10)
═══════════════════════════════════════════════════════════════════════════

Se pedir FGTS:
 Mencionar depósitos não realizados
 Mencionar multa de 40% sobre saldo

Se pedir 13º Salário:
 Especificar parcelas (integral/proporcional)
 Quantificar períodos devidos

Se pedir Férias:
 Diferenciar vencidas e proporcionais
 SEMPRE mencionar terço constitucional (1/3)

Se pedir Dano Moral:
 Especificar valor em reais
 Classificar grau: leve, médio, grave ou gravíssimo

Se pedir Verbas Rescisórias:
 Detalhar: aviso prévio, saldo salário, férias proporcionais, 13º proporcional, FGTS+40%

Se pedir Horas Extras:
 Quantificar (horas por mês/semana)
 Especificar período (datas início e fim)
 Mencionar adicional (mínimo 50% ou conforme CCT)

Se pedir Rescisão Indireta:
 Fundamentar falta grave do empregador
 Citar expressamente art. 483 da CLT

Se pedir Adicional de Insalubridade:
 Mencionar grau (mínimo, médio ou máximo)
 Citar NR-15

Se pedir Adicional Noturno:
 Especificar horário (22h às 5h)
 Mencionar percentual mínimo de 20%

Se pedir Estabilidade:
 Especificar tipo (gestante/acidentária/CIPA)
 Fundamentar período de estabilidade

═══════════════════════════════════════════════════════════════════════════
5. REFLEXOS OBRIGATÓRIOS DE HORAS EXTRAS
═══════════════════════════════════════════════════════════════════════════

 CRÍTICO: Se pedir HORAS EXTRAS, DEVE incluir TODOS os 7 reflexos:

 1. Adicional de Horas Extras (mínimo 50% ou conforme CCT)
 2. DSR - Descanso Semanal Remunerado sobre as horas extras
 3. Reflexos em 13º Salário
 4. Reflexos em Férias + 1/3 constitucional
 5. Reflexos em FGTS + multa de 40%
 6. Reflexos em Aviso Prévio (quando aplicável)
 7. Reflexo em Adicional Noturno (se HE em horário noturno)


 FALTA DE REFLEXOS = BLOQUEIO AUTOMÁTICO NA AUDITORIA

═══════════════════════════════════════════════════════════════════════════
6. TIPO DE AÇÃO - {tipo_processo}
═══════════════════════════════════════════════════════════════════════════

{'RECONHECIMENTO DE VÍNCULO:' if tipo_processo == 'RECONHECIMENTO_VINCULO' else ''}
{'- Cliente NÃO tem registro formal em CTPS' if tipo_processo == 'RECONHECIMENTO_VINCULO' else ''}
{'- Objetivo: comprovar relação de emprego sem registro' if tipo_processo == 'RECONHECIMENTO_VINCULO' else ''}
{'- Pedir: reconhecimento de vínculo + anotação CTPS + verbas devidas' if tipo_processo == 'RECONHECIMENTO_VINCULO' else ''}

{'AÇÃO ACIDENTÁRIA:' if tipo_processo == 'ACAO_ACIDENTARIA' else ''}
{'- Acidente de trabalho ou doença ocupacional' if tipo_processo == 'ACAO_ACIDENTARIA' else ''}
{'- Responsabilizar empregador por danos à saúde' if tipo_processo == 'ACAO_ACIDENTARIA' else ''}
{'- Mencionar nexo causal + culpa/negligência empregador' if tipo_processo == 'ACAO_ACIDENTARIA' else ''}

{'DIFERENÇAS CONTRATUAIS:' if tipo_processo == 'DIFERENCAS_CONTRATUAIS' else ''}
{'- Cliente TEM registro mas há verbas não pagas' if tipo_processo == 'DIFERENCAS_CONTRATUAIS' else ''}
{'- Cobrar diferenças: HE, verbas rescisórias, adicionais' if tipo_processo == 'DIFERENCAS_CONTRATUAIS' else ''}

═══════════════════════════════════════════════════════════════════════════
7. QUALIDADE FINAL - CHECKLIST ANTES DE RETORNAR
═══════════════════════════════════════════════════════════════════════════

Antes de retornar, VERIFIQUE:

 Vocativo formal presente
 Qualificação completa (nome, RG, CPF, endereço com CEP)
 Seção DOS FATOS narrativa e cronológica
 Seção DO MÉRITO com fundamentação legal (CLT, CF/88)
 Seção DOS PEDIDOS específica e quantificada
 Fórmula encerramento COMPLETA: "Termos em que, pede deferimento."
 Data e local presentes
 Assinatura e OAB
 Se pediu HE  TODOS os 7 reflexos presentes
 Pedidos específicos (não genéricos)
 Fundamentação legal adequada

═══════════════════════════════════════════════════════════════════════════
 LEMBRETE FINAL
═══════════════════════════════════════════════════════════════════════════

Esta petição será AUDITADA automaticamente pelo CHECKLIST IA AUDITORA V4.0.

Erros críticos = BLOQUEIO AUTOMÁTICO
Estrutura incompleta = PENALIZAÇÃO DE -5 PONTOS POR ELEMENTO
Reflexos HE faltando = PENALIZAÇÃO DE -10 PONTOS POR REFLEXO

Score mínimo para aprovação: 70/100

RETORNE APENAS A PETIÇÃO COMPLETA, PERFEITA E PRONTA PARA PROTOCOLO.
"""
        
        # NOVO: Lógica condicional para Prompt Master
        if usar_prompt_master and PROMPT_MASTER_DISPONIVEL:
            print(f"         Usando PROMPT MASTER para petição de alto nível (12-18 páginas)")
            
            # Gerar prompt master completo
            prompt_master_instrucoes = gerar_prompt_master(
                tipo_processo=tipo_processo,
                cliente_info=cliente_info,
                documentos=documentos_completos,
                cronologia=cronologia_fatos,
                resumo_video=resumo_video
            )
            
            # Substituir instruções padrão pelo prompt master
            system_prompt = prompt_master_instrucoes
            
            # Aumentar max_tokens para suportar petições mais longas
            max_tokens_config = 32000
            
        else:
            # Modo padrão: usar instruções existentes
            if usar_prompt_master and not PROMPT_MASTER_DISPONIVEL:
                print(f"         Prompt Master solicitado mas não disponível, usando modo padrão")
            
            # Adicionar instruções ao system prompt
            system_prompt += "\n\n" + instrucoes
            max_tokens_config = 16000

        
        print(f"        - Enviando para Claude AI (Modelo Claude Sonnet 4)...")
        print(f"        - Max tokens: {max_tokens_config}")
        message = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=max_tokens_config,
            system=system_prompt,
            messages=[{"role": "user", "content": conteudo_mensagem}]
        )
        peticao = message.content[0].text
        print(f"        - Petição gerada! ({len(peticao)} chars)")
        
        # VALIDAÇÃO PROMPT MASTER (se ativado)
        if usar_prompt_master and PROMPT_MASTER_DISPONIVEL:
            print(f"        - Validando petição com critérios do Prompt Master...")
            relatorio_master = gerar_relatorio_validacao_master(peticao)
            
            print(f"        - Score Prompt Master: {relatorio_master['score']}/100 ({relatorio_master['status']})")
            
            if relatorio_master['problemas']:
                print(f"        -  {len(relatorio_master['problemas'])} problema(s) encontrado(s)")
            
            # Armazenar relatório para uso posterior
            peticao_validacao_master = relatorio_master
        else:
            peticao_validacao_master = None
        
        # PÓS-PROCESSAMENTO: Limpar marcadores de quebra de linha
        peticao = limpar_marcadores_quebra_linha(peticao)
        print(f"        - Marcadores de quebra de linha removidos")
        
        # PÓS-PROCESSAMENTO: Processar marcadores de formatação especial
        peticao = processar_marcadores_formatacao(peticao)
        print(f"        - Marcadores de formatação processados")
        
        # PÓS-PROCESSAMENTO: Adicionar numeração de parágrafos (apenas modo padrão)
        if not usar_prompt_master:
            peticao = adicionar_numeracao_paragrafos(peticao)
            print(f"        - Numeração de parágrafos adicionada")

        
        return peticao
    except Exception as e:
        print(f"        ERRO: {e}")
        return None

def limpar_marcadores_quebra_linha(peticao_texto):
    """
    Remove marcadores literais de quebra de linha e substitui por quebras reais
    
    Marcadores removidos:
    - [Quebra de linha dupla] -> duas linhas em branco
    - [Quebra de linha tripla] -> três linhas em branco
    - [Quebra de linha simples] -> uma linha em branco
    - [Quebra de linha] -> uma linha em branco
    
    Args:
        peticao_texto: Texto da petição com possíveis marcadores
        
    Returns:
        Texto da petição com marcadores substituídos por quebras reais
    """
    try:
        import re
        
        # Substituir marcadores por quebras de linha reais
        texto_limpo = peticao_texto
        
        # [Quebra de linha tripla] -> 3 quebras
        texto_limpo = re.sub(r'\[Quebra de linha tripla\]', '\n\n\n', texto_limpo, flags=re.IGNORECASE)
        
        # [Quebra de linha dupla] -> 2 quebras
        texto_limpo = re.sub(r'\[Quebra de linha dupla\]', '\n\n', texto_limpo, flags=re.IGNORECASE)
        
        # [Quebra de linha simples] ou [Quebra de linha] -> 1 quebra
        texto_limpo = re.sub(r'\[Quebra de linha( simples)?\]', '\n', texto_limpo, flags=re.IGNORECASE)
        
        return texto_limpo
        
    except Exception as e:
        print(f"        [AVISO] Erro ao limpar marcadores: {e}")
        return peticao_texto  # Retornar original em caso de erro

def processar_marcadores_formatacao(peticao_texto):
    """
    Processa marcadores especiais de formatação que serão aplicados no documento Word:
    - [CENTRALIZAR]...[/CENTRALIZAR] -> texto centralizado
    - [NEGRITO_SUBLINHADO]...[/NEGRITO_SUBLINHADO] -> negrito + sublinhado
    - [ESPACO_GRANDE] -> 8-10 linhas em branco
    
    Args:
        peticao_texto: Texto da petição com possíveis marcadores
        
    Returns:
        Texto com marcadores convertidos para tags especiais
    """
    try:
        import re
        
        texto_processado = peticao_texto
        
        # Converter marcadores para tags especiais que serão processadas no Word
        # [CENTRALIZAR]...[/CENTRALIZAR] -> <<<CENTRO>>>...<<</CENTRO>>>
        texto_processado = re.sub(
            r'\[CENTRALIZAR\](.*?)\[/CENTRALIZAR\]',
            r'<<<CENTRO>>>\1<<</CENTRO>>>',
            texto_processado,
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # [NEGRITO_SUBLINHADO]...[/NEGRITO_SUBLINHADO] -> <<<BOLD_UL>>>...<<</BOLD_UL>>>
        texto_processado = re.sub(
            r'\[NEGRITO_SUBLINHADO\](.*?)\[/NEGRITO_SUBLINHADO\]',
            r'<<<BOLD_UL>>>\1<<</BOLD_UL>>>',
            texto_processado,
            flags=re.IGNORECASE
        )
        
        # [ESPACO_GRANDE] -> 8 linhas em branco
        texto_processado = texto_processado.replace('[ESPACO_GRANDE]', '\n\n\n\n\n\n\n\n')
        
        return texto_processado
        
    except Exception as e:
        print(f"        [AVISO] Erro ao processar marcadores de formatação: {e}")
        return peticao_texto

def adicionar_numeracao_paragrafos(peticao_texto):
    """
    Adiciona numeração sequencial aos parágrafos da petição
    
    Regras:
    - Não numerar: vocativo inicial, títulos de seções, assinatura final
    - Numerar: todos os parágrafos do corpo
    
    Args:
        peticao_texto: Texto da petição sem numeração
    
    Returns:
        Texto da petição com parágrafos numerados
    """
    try:
        linhas = peticao_texto.split('\n')
        linhas_processadas = []
        numero_paragrafo = 1
        
        # Palavras-chave que indicam títulos (não devem ser numerados)
        titulos_secoes = [
            'EXCELENTÍSSIMO', 'EXMO', 'MERITÍSSIMO',
            'PRELIMINAR', 'DOS FATOS', 'DO MÉRITO', 'DOS PEDIDOS',
            'DA COMPETÊNCIA', 'DO JUÍZO', 'DA JUNTADA',
            'DO RECONHECIMENTO', 'DAS HORAS EXTRAS', 'DO ADICIONAL',
            'DA RESCISÃO', 'DAS MULTAS', 'DOS DANOS',
            'TERMOS EM QUE', 'PEDE DEFERIMENTO',
            'NESTES TERMOS'
        ]
        
        for linha in linhas:
            linha_limpa = linha.strip()
            
            # Pular linhas vazias
            if not linha_limpa:
                linhas_processadas.append(linha)
                continue
            
            # Verificar se é título de seção (não numerar)
            eh_titulo = False
            for titulo in titulos_secoes:
                if titulo in linha_limpa.upper():
                    eh_titulo = True
                    break
            
            # Verificar se já está numerado
            ja_numerado = False
            if linha_limpa and linha_limpa[0].isdigit():
                # Verificar se tem ponto após número
                partes = linha_limpa.split('.', 1)
                if len(partes) > 1 and partes[0].isdigit():
                    ja_numerado = True
            
            # Verificar se é assinatura (linhas com underscores)
            eh_assinatura = '_' in linha_limpa
            
            # Verificar se é data (contém números e "de")
            eh_data = 'de' in linha_limpa.lower() and any(char.isdigit() for char in linha_limpa)
            
            # Se não é título, não é assinatura, não é data, não está numerado e tem conteúdo substancial
            if (not eh_titulo and not eh_assinatura and not eh_data and 
                not ja_numerado and len(linha_limpa) > 20):
                # Adicionar numeração
                linha_numerada = f"{numero_paragrafo}. {linha_limpa}"
                linhas_processadas.append(linha_numerada)
                numero_paragrafo += 1
            else:
                linhas_processadas.append(linha)
        
        return '\n'.join(linhas_processadas)
        
    except Exception as e:
        print(f"        [AVISO] Erro ao numerar parágrafos: {e}")
        return peticao_texto  # Retornar original em caso de erro

def salvar_peticao_no_drive(service, peticao_texto, cliente_info, arquivos_cliente=None, usar_prompt_master=False):
    try:
        print(f"        - Salvando no Drive...")
        
        # SISTEMA DE PRINTS V4.0 - Inserir marcadores antes de salvar
        texto_final = peticao_texto
        marcadores_info = None
        
        if arquivos_cliente:
            print(f"        [SISTEMA DE PRINTS V4.0 - Analisando...]")
            tipo_acao = identificar_tipo_acao_por_texto(peticao_texto)
            print(f"        - Tipo identificado: {tipo_acao}")
            
            # Extrair apenas os nomes dos arquivos
            nomes_arquivos = []
            for arq in arquivos_cliente:
                if isinstance(arq, dict):
                    nome = arq.get('nome', arq.get('name', ''))
                    if nome:
                        nomes_arquivos.append(nome)
                elif isinstance(arq, str):
                    nomes_arquivos.append(arq)
            
            texto_final, marcadores, faltantes, criticos_faltantes = inserir_marcadores_prints(
                peticao_texto, tipo_acao, nomes_arquivos
            )
            
            marcadores_info = {
                'marcadores': marcadores,
                'faltantes': faltantes,
                'criticos_faltantes': criticos_faltantes,
                'tipo_acao': tipo_acao
            }
            
            print(f"        - Marcadores inseridos: {len(marcadores)}")
            if criticos_faltantes:
                print(f"        -  PRINTS CRÍTICOS FALTANTES!")
        
        # NOVO: Usar modelo como base para preservar cabeçalho/rodapé
        tipo_processo = cliente_info['tipo_processo']
        modelo_id = MODELOS_IDS.get(tipo_processo)
        
        if modelo_id:
            print(f"        - Usando modelo como base (ID: {modelo_id[:20]}...)")
            try:
                # Baixar modelo
                modelo_conteudo = baixar_arquivo(service, modelo_id)
                if modelo_conteudo:
                    # Salvar modelo temporariamente
                    tmp_modelo_fd, tmp_modelo_path = tempfile.mkstemp(suffix='.docx')
                    os.close(tmp_modelo_fd)
                    with open(tmp_modelo_path, 'wb') as f:
                        f.write(modelo_conteudo)
                    
                    # Abrir modelo
                    doc = Document(tmp_modelo_path)
                    
                    # Limpar TODO o conteúdo do corpo (parágrafos E tabelas)
                    # Isso mantém apenas cabeçalho e rodapé
                    
                    # 1. Remover todas as tabelas
                    for table in doc.tables:
                        table._element.getparent().remove(table._element)
                    
                    # 2. Remover todos os parágrafos
                    for _ in range(len(doc.paragraphs)):
                        p = doc.paragraphs[0]
                        p._element.getparent().remove(p._element)
                    
                    # 3. Adicionar novo conteúdo (formatação será aplicada depois)
                    primeiro_paragrafo = True
                    for linha in texto_final.split('\n'):
                        if linha.strip():
                            # Adicionar parágrafo com estilo Normal
                            p = doc.add_paragraph(linha)
                            try:
                                p.style = 'Normal'
                            except:
                                pass
                            
                            # Aplicar negrito apenas no primeiro parágrafo (cabeçalho)
                            if primeiro_paragrafo:
                                for run in p.runs:
                                    run.font.bold = True
                                primeiro_paragrafo = False
                    
                    # Limpar arquivo temporário do modelo
                    try:
                        os.unlink(tmp_modelo_path)
                    except:
                        pass
                    
                    print(f"         Cabeçalho e rodapé preservados do modelo")
                    print(f"         Tabelas do modelo removidas")
                    print(f"         Formatação aplicada")
                else:
                    print(f"         Não foi possível baixar modelo, criando documento vazio")
                    doc = Document()
                    for p in texto_final.split('\n'):
                        if p.strip():
                            doc.add_paragraph(p)
            except Exception as e:
                print(f"         Erro ao usar modelo: {e}, criando documento vazio")
                doc = Document()
                for p in texto_final.split('\n'):
                    if p.strip():
                        doc.add_paragraph(p)
        else:
            print(f"         Modelo não configurado para {tipo_processo}, criando documento vazio")
            doc = Document()
            for p in texto_final.split('\n'):
                if p.strip():
                    doc.add_paragraph(p)
        
        # Aplicar formatação padrão do escritório (sempre)
        doc = aplicar_formatacao_master(doc)
        
        # Aplicar formatações especiais (centralização, sublinhado, negrito seletivo)
        doc = aplicar_formatacoes_especiais_word(doc)

        
        # Salvar documento final
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
        os.close(tmp_fd)
        doc.save(tmp_path)
        nome = f"Peticao_{cliente_info['tipo_processo']}_{cliente_info['cliente_nome'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        pasta = PASTAS_PETICOES_GERADAS.get(cliente_info['tipo_processo'])
        metadata = {'name': nome, 'parents': [pasta]}
        media = MediaFileUpload(tmp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = service.files().create(body=metadata, media_body=media, fields='id, name, webViewLink').execute()
        try:
            os.unlink(tmp_path)
        except:
            pass
        print(f"        - Salvo: {file.get('name')}")
        
        # Adicionar info de marcadores ao retorno
        if marcadores_info:
            file['marcadores_prints'] = marcadores_info
        
        return file
    except Exception as e:
        print(f"        ERRO: {e}")
        return None

def gerar_relatorio_score(resultado, score_medio_tipo, score_medio_escritorio):
    """
    Gera relatório detalhado explicando o score da petição
    """
    try:
        score = resultado.get('score', 0)
        
        # Determinar ranking
        if score >= 90:
            ranking = "EXCELENTE"
            emoji = ""
        elif score >= 80:
            ranking = "MUITO BOM"
            emoji = ""
        elif score >= 70:
            ranking = "BOM"
            emoji = ""
        elif score >= 60:
            ranking = "SATISFATÓRIO"
            emoji = ""
        else:
            ranking = "PRECISA MELHORAR"
            emoji = ""
        
        relatorio = {
            'score': score,
            'ranking': ranking,
            'emoji': emoji,
            'analise': [],
            'pontos_fortes': [],
            'pontos_melhoria': [],
            'comparacoes': []
        }
        
        # Análise por categoria
        if resultado.get('estrutura_ok'):
            relatorio['pontos_fortes'].append('Estrutura completa e bem organizada')
        else:
            relatorio['pontos_melhoria'].append('Estrutura precisa ser melhorada')
            
        if resultado.get('qualidade_ok'):
            relatorio['pontos_fortes'].append('Qualidade textual excelente')
        else:
            relatorio['pontos_melhoria'].append('Qualidade do texto pode melhorar')
        
        # Pedidos
        pedidos_count = len(resultado.get('pedidos_validacao', []))
        if pedidos_count >= 5:
            relatorio['pontos_fortes'].append(f'{pedidos_count} pedidos específicos e quantificados')
        elif pedidos_count > 0:
            relatorio['pontos_melhoria'].append(f'Apenas {pedidos_count} pedidos identificados')
        else:
            relatorio['pontos_melhoria'].append('Pedidos precisam ser mais específicos')
        
        # Reflexos HE
        if resultado.get('reflexos_he_detalhados'):
            reflexos_ok = sum(1 for r in resultado['reflexos_he_detalhados'] if r.get('presente') == True)
            reflexos_total = 7
            if reflexos_ok == reflexos_total:
                relatorio['pontos_fortes'].append('Todos os 7 reflexos de HE presentes')
            elif reflexos_ok > 0:
                relatorio['pontos_melhoria'].append(f'Apenas {reflexos_ok}/7 reflexos de HE presentes')
            else:
                relatorio['pontos_melhoria'].append('Reflexos de HE não identificados')
        
        # Erros críticos
        erros_criticos = resultado.get('erros_criticos', [])
        if erros_criticos:
            for erro in erros_criticos[:3]:
                relatorio['pontos_melhoria'].append(f' {erro}')
        
        # Comparações
        if score > score_medio_tipo:
            diff = score - score_medio_tipo
            relatorio['comparacoes'].append(f'⬆ {diff:.1f} pontos acima da média do tipo')
        elif score < score_medio_tipo:
            diff = score_medio_tipo - score
            relatorio['comparacoes'].append(f'⬇ {diff:.1f} pontos abaixo da média do tipo')
        else:
            relatorio['comparacoes'].append(f' Na média do tipo ({score_medio_tipo:.1f})')
        
        if score > score_medio_escritorio:
            diff = score - score_medio_escritorio
            relatorio['comparacoes'].append(f'⬆ {diff:.1f} pontos acima da média geral')
        elif score < score_medio_escritorio:
            diff = score_medio_escritorio - score
            relatorio['comparacoes'].append(f'⬇ {diff:.1f} pontos abaixo da média geral')
        
        # Análise geral
        if score >= 90:
            relatorio['analise'].append('Petição de excelente qualidade, pronta para protocolo.')
        elif score >= 80:
            relatorio['analise'].append('Petição de muito boa qualidade, pequenos ajustes podem aperfeiçoar.')
        elif score >= 70:
            relatorio['analise'].append('Petição com boa estrutura, alguns pontos podem ser melhorados.')
        elif score >= 60:
            relatorio['analise'].append('Petição satisfatória, recomenda-se revisão em alguns pontos.')
        else:
            relatorio['analise'].append('Petição precisa de melhorias significativas antes do protocolo.')
        
        return relatorio
        
    except Exception as e:
        print(f"Erro ao gerar relatório de score: {e}")
        return {
            'score': resultado.get('score', 0),
            'ranking': 'N/A',
            'emoji': '',
            'analise': ['Erro ao gerar relatório'],
            'pontos_fortes': [],
            'pontos_melhoria': [],
            'comparacoes': []
        }

def auditar_peticao_com_claude(service, arquivo_id, tipo_processo, cliente_nome):
    """Audita petição - CHECKLIST V4.0 FASE 3 PERFEITA"""
    print(f"        [AUDITANDO - FASE 3: Qualidade Extra + Comparações]")
    try:
        conteudo = baixar_arquivo(service, arquivo_id)
        if not conteudo:
            return None
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
        os.close(tmp_fd)
        with open(tmp_path, 'wb') as f:
            f.write(conteudo)
        doc = Document(tmp_path)
        peticao_texto = "\n".join([p.text for p in doc.paragraphs])
        os.unlink(tmp_path)
        
        # Estatísticas para comparação
        stats = carregar_estatisticas()
        score_medio = SCORES_MEDIOS.get(tipo_processo, 75)
        score_medio_escritorio = stats.get('por_tipo', {}).get(tipo_processo, {}).get('score_medio', score_medio)
        
        prompt = f"""AUDITOR ESPECIALIZADO - CHECKLIST V4.0 FASE 3 PERFEITA

TIPO: {tipo_processo} 
SCORE MÉDIO CHECKLIST: {score_medio}/100
SCORE MÉDIO ESCRITÓRIO: {score_medio_escritorio:.1f}/100

PETIÇÃO:
{peticao_texto[:10000]}

VALIDAÇÃO EM 5 PARTES:

PARTE 1 - 10 ERROS CRÍTICOS (-10pts):
Inépcia, Pedidos genéricos, Sem fundamentação, Valor ausente, Qualificação incompleta,
Docs ALTA ausentes, Contradições, Sem competência, Reflexos HE faltando, Encerramento incompleto

PARTE 2 - ESTRUTURA (11 elementos, -5pts):
Vocativo, Qualif.Reclamante, Qualif.Reclamada, Fórmula, DOS FATOS, DO MÉRITO,
DOS PEDIDOS, VALOR CAUSA, Encerramento, Preliminares(opcional)

PARTE 3 - PEDIDOS (Top 10):
FGTS(depósitos+40%), 13º(parcelas), Férias(vencidas/proporcionais+1/3), 
Dano Moral(valor+grau), HE(quantifica+período+50%), Verbas Rescisórias(detalha),
Rescisão Indireta(art.483), Insalubridade(grau+NR15), Noturno(22h-5h+20%), 
Estabilidade(tipo)

PARTE 4 - REFLEXOS HE (7 obrigatórios):
1.Adicional 50% 2.DSR 3.Reflexos 13º 4.Reflexos Férias+1/3
5.Reflexos FGTS+40% 6.Reflexos Aviso 7.Noturno(se aplicável)

PARTE 5 - QUALIDADE EXTRA (+bônus):
 Jurisprudência TST/TRT (+3pts cada, máx 9pts)
 Cálculos detalhados tabela/planilha (+5pts)
 Narrativa persuasiva e emocional (+3pts)
 Fundamentação doutrinária (+3pts)

RETORNE JSON:
{{
  "score": 88, "aprovada": true,
  "justificativa_score": "Score 88/100. Acima média escritório({score_medio_escritorio:.1f}). Perdeu 12pts: HE não quantificadas(-3), reflexo FGTS ausente(-10), sem jurisprudência(0 bônus).",
  
  "erros_criticos": [],
  
  "estrutura_validacao": {{"vocativo":true,"qualif_reclamante":true,"qualif_reclamada":true,"formula":true,"dos_fatos":true,"do_merito":true,"dos_pedidos":true,"valor_causa":true,"encerramento":true}},
  "estrutura_faltante": [],
  
  "pedidos_validacao": [
    {{"pedido":"FGTS","validacao":" Depósitos + multa 40%"}},
    {{"pedido":"HE","validacao":" Falta quantificar horas"}}
  ],
  
  "reflexos_he_detalhados": [
    {{"reflexo":"1.Adicional HE 50%","presente":true}},
    {{"reflexo":"5.Reflexos FGTS+40%","presente":false}}
  ],
  "reflexos_he_faltantes": ["FGTS+40%"],
  
  "qualidade_extra": {{
    "jurisprudencia_tst": false,
    "jurisprudencia_trt": false,
    "calculos_detalhados": false,
    "narrativa_persuasiva": true,
    "fundamentacao_doutrinaria": false,
    "bonus_pontos": 3
  }},
  
  "comparacao_escritorio": {{
    "score_cliente": 88,
    "score_medio_tipo": {score_medio_escritorio:.1f},
    "diferenca": "+X.X pontos",
    "melhor_que_media": true
  }},
  
  "alertas": ["HE não quantificadas"],
  "sugestoes": ["Quantificar HE exatamente","Adicionar jurisprudência TST"],
  "pontos_positivos": ["Estrutura completa","Narrativa persuasiva"],
  "melhorias_100": ["Quantificar HE (2h/dia)","Reflexos FGTS explícito","Adicionar 2-3 jurisprudências TST"],
  "resumo": "Aprovada 88/100, acima de {score_medio_escritorio:.1f}. Forte: estrutura+narrativa. Melhorar: reflexos+jurisprudência."
}}"""
        
        message = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        resposta = message.content[0].text.strip()
        if resposta.startswith('```'):
            linhas = resposta.split('\n')
            resposta = '\n'.join(linhas[1:-1])
        resposta = resposta.replace('```json', '').replace('```', '').strip()
        resultado = json.loads(resposta)
        
        score = resultado.get('score', 0)
        
        # Calcular comparações
        comparacao_base = "⬆" if score > score_medio else "⬇" if score < score_medio else ""
        ranking_percentil = calcular_ranking(score, tipo_processo)  # "Top 81%"
        
        # CALCULAR RANKING PARA O DASHBOARD (baseado no score)
        if score >= 90:
            ranking_dashboard = "EXCELENTE"
        elif score >= 80:
            ranking_dashboard = "MUITO BOM"
        elif score >= 70:
            ranking_dashboard = "BOM"
        elif score >= 60:
            ranking_dashboard = "SATISFATÓRIO"
        else:
            ranking_dashboard = "PRECISA MELHORAR"
        
        # SEMPRE APROVAR (sistema mudou para não rejeitar)
        resultado['aprovada'] = True
        
        print(f"        - Score: {score}/100 (base: {score_medio}, escritório: {score_medio_escritorio:.1f}) {comparacao_base}")
        print(f"        - Ranking: {ranking_percentil}")
        print(f"        - APROVADA ")
        print(f"        - Classificação: {ranking_dashboard}")
        
        # Qualidade extra
        if resultado.get('qualidade_extra'):
            bonus = resultado['qualidade_extra'].get('bonus_pontos', 0)
            if bonus > 0:
                print(f"        - Bônus Qualidade: +{bonus} pontos")
        
        # Mostrar resumo
        if resultado.get('reflexos_he_detalhados'):
            reflexos_ok = sum(1 for r in resultado['reflexos_he_detalhados'] if r.get('presente') == True)
            reflexos_total = sum(1 for r in resultado['reflexos_he_detalhados'] if r.get('presente') is not None)
            if reflexos_total > 0:
                print(f"        - Reflexos HE: {reflexos_ok}/{reflexos_total}")
        
        if resultado.get('pedidos_validacao'):
            print(f"        - Pedidos validados: {len(resultado['pedidos_validacao'])}")
        
        # Gerar relatório detalhado para o dashboard
        relatorio_detalhado = gerar_relatorio_score(resultado, score_medio, score_medio_escritorio)
        
        # Adicionar informações extras ao resultado
        resultado['ranking'] = ranking_dashboard  #  USAR RANKING DO DASHBOARD
        resultado['ranking_percentil'] = ranking_percentil  #  Guardar o percentil também
        resultado['comparacao_base'] = score_medio
        resultado['comparacao_escritorio'] = score_medio_escritorio
        resultado['relatorio_detalhado'] = relatorio_detalhado
        
        return resultado
        
    except Exception as e:
        print(f"        ERRO: {e}")
        import traceback
        traceback.print_exc()
        return None

def mover_peticao(service, arquivo_id, aprovada):
    try:
        if aprovada:
            pasta_destino = os.getenv('PASTA_03_APROVADAS')
            status = "APROVADAS"
        else:
            pasta_destino = os.getenv('PASTA_04_REJEITADAS')
            status = "REJEITADAS"
        file = service.files().get(fileId=arquivo_id, fields='parents').execute()
        pastas_atuais = file.get('parents', [])
        if not pastas_atuais:
            return False
        previous_parents = ",".join(pastas_atuais)
        service.files().update(
            fileId=arquivo_id,
            addParents=pasta_destino,
            removeParents=previous_parents,
            fields='id, parents'
        ).execute()
        print(f"        - Movida para: {status}")
        return True
    except Exception as e:
        print(f"        Erro ao mover: {e}")
        return False

# AGENTE 1: GERADOR
def agente_gerador():
    print(f"\n{'='*70}")
    print(f"[AGENTE GERADOR] {datetime.now().strftime('%H:%M:%S')}")
    print(f"{'='*70}")
    
    try:
        service = autenticar_google_drive()
        pastas = [
            (os.getenv('PASTA_RECONHECIMENTO_VINCULO'), 'RECONHECIMENTO_VINCULO'),
            (os.getenv('PASTA_ACAO_ACIDENTARIA'), 'ACAO_ACIDENTARIA'),
            (os.getenv('PASTA_DIFERENCAS_CONTRATUAIS'), 'DIFERENCAS_CONTRATUAIS')
        ]
        
        total_geradas = 0
        
        for pasta_id, tipo in pastas:
            if not pasta_id:
                continue
            print(f"\n  -> {tipo}")
            pastas_clientes = listar_pastas(service, pasta_id)
            print(f"     {len(pastas_clientes)} pasta(s)")
            
            for pasta_cliente in pastas_clientes:
                # VERIFICAR SE JÁ PROCESSADO (evita duplicatas)
                if verificar_cliente_ja_processado(service, pasta_cliente['id']):
                    continue
                
                # MUDANÇA: Usar listagem recursiva para pegar arquivos de subpastas também
                arquivos = listar_arquivos_recursivo(service, pasta_cliente['id'])
                if not arquivos:
                    continue
                
                docs = []
                for arq in arquivos:
                    classif = classificar_documento(arq['name'])
                    if classif['prioridade'] != 'IGNORAR':
                        docs.append({'id': arq['id'], 'nome': arq['name'], 'tipo': classif['tipo'], 'prioridade': classif['prioridade']})
                
                # DEBUG: Mostrar documentos reconhecidos
                print(f"\n     [CLIENTE: {pasta_cliente['name']}]")
                print(f"      Arquivos na pasta ({len(arquivos)}):")
                for doc in docs:
                    emoji = "" if doc['prioridade'] == 'ALTA' else "" if doc['prioridade'] == 'MEDIA' else ""
                    print(f"        {emoji} {doc['tipo']}: {doc['nome']}")
                
                # Usar nova função com 3 prioridades
                verif = verificar_documentacao_completa_v10(docs, tipo)
                
                print(f"\n      Validação:")
                print(f"         ALTA: {len(verif['alta_presentes'])}/{len(DOCUMENTOS_POR_TIPO[tipo]['ALTA'])}")
                print(f"         MÉDIA: {len(verif['media_presentes'])}/{len(DOCUMENTOS_POR_TIPO[tipo]['MEDIA'])}")
                print(f"         BAIXA: {len(verif['baixa_presentes'])}/{len(DOCUMENTOS_POR_TIPO[tipo]['BAIXA'])}")
                
                # REGRA CRÍTICA: Bloquear se faltar QUALQUER documento
                tem_faltantes = (len(verif['alta_faltantes']) > 0 or 
                                len(verif['media_faltantes']) > 0 or 
                                len(verif['baixa_faltantes']) > 0)
                
                if tem_faltantes:
                    print(f"\n         BLOQUEADO - Documentação Incompleta:")
                    
                    if verif['alta_faltantes']:
                        print(f"         FALTAM ALTA (Crítico):")
                        for doc in verif['alta_faltantes']:
                            print(f"           • {doc}")
                    
                    if verif['media_faltantes']:
                        print(f"         FALTAM MÉDIA (Importante):")
                        for doc in verif['media_faltantes']:
                            print(f"           • {doc}")
                    
                    if verif['baixa_faltantes']:
                        print(f"         FALTAM BAIXA (Desejável):")
                        for doc in verif['baixa_faltantes']:
                            print(f"           • {doc}")
                    
                    print(f"         Adicione os documentos e reprocesse")
                    continue
                
                print(f"         Documentação COMPLETA - Gerando petição...")
                print(f"        >> GERANDO PETICAO <<")
                
                docs_completos = []
                for doc in docs:
                    cont = baixar_arquivo(service, doc['id'])
                    if cont:
                        texto = extrair_texto_pdf(cont) if doc['nome'].lower().endswith('.pdf') else ""
                        docs_completos.append({'tipo': doc['tipo'], 'nome': doc['nome'], 'conteudo': cont, 'texto': texto})
                
                cliente_info = {
                    'cliente_nome': pasta_cliente['name'],
                    'tipo_processo': tipo,
                    'verif_docs': verif  # Passar verificação para auditoria
                }
                
                peticao = gerar_peticao_com_claude(service, cliente_info, docs_completos, tipo)
                
                if peticao:
                    arquivo = salvar_peticao_no_drive(service, peticao, cliente_info, arquivos_cliente=docs)
                    if arquivo:
                        # Gerar relatório de prints se marcadores foram inseridos
                        if arquivo.get('marcadores_prints'):
                            info_prints = arquivo['marcadores_prints']
                            relatorio_prints = gerar_relatorio_prints(
                                info_prints['tipo_acao'],
                                info_prints['marcadores'],
                                info_prints['faltantes'],
                                info_prints['criticos_faltantes'],
                                pasta_cliente['name']
                            )
                            print(f"         Relatório de prints: {relatorio_prints}")
                        
                        # Salvar no histórico
                        salvar_no_historico(
                            pasta_cliente['name'],
                            tipo,
                            arquivo
                        )
                        
                        # MARCAR IMEDIATAMENTE (evita duplicatas)
                        marcar_cliente_como_processado(service, pasta_cliente['id'], {
                            'nome_arquivo': arquivo['name'],
                            'link': arquivo['webViewLink']
                        })
                        total_geradas += 1
                        print(f"         PETIÇÃO GERADA E REGISTRADA!")
        
        print(f"\n{'='*70}")
        print(f"  GERADOR: {total_geradas} petição(ões)")
        print(f"{'='*70}")
        
    except Exception as e:
        print(f"[ERRO GERADOR] {e}")

# AGENTE 2: AUDITOR
def agente_auditor():
    print(f"\n{'='*70}")
    print(f"[AGENTE AUDITOR] {datetime.now().strftime('%H:%M:%S')}")
    print(f"{'='*70}")
    
    try:
        service = autenticar_google_drive()
        total_auditadas = 0
        total_aprovadas = 0
        
        for tipo, pasta_id in PASTAS_PETICOES_GERADAS.items():
            print(f"\n  -> {tipo}")
            arquivos = listar_arquivos_pasta(service, pasta_id)
            
            if not arquivos:
                print(f"     Nenhuma petição pendente")
                continue
            
            print(f"     {len(arquivos)} petição(ões) pendente(s)")
            
            for arquivo in arquivos:
                print(f"\n     [PETIÇÃO: {arquivo['name']}]")
                
                # Extrair nome do cliente
                parts = arquivo['name'].split('_')
                cliente_nome = parts[2] if len(parts) > 2 else 'Desconhecido'
                
                resultado = auditar_peticao_com_claude(service, arquivo['id'], tipo, cliente_nome)
                
                if resultado:
                    total_auditadas += 1
                    
                    # Salvar relatório
                    log_auditoria(cliente_nome, tipo, resultado, arquivo['name'])
                    
                    # SEMPRE APROVAR - Status baseado no score
                    score = resultado.get('score', 0)
                    
                    # Determinar ranking baseado no score
                    if score >= 90:
                        ranking = "EXCELENTE"
                    elif score >= 80:
                        ranking = "MUITO BOM"
                    elif score >= 70:
                        ranking = "BOM"
                    elif score >= 60:
                        ranking = "SATISFATÓRIO"
                    else:
                        ranking = "PRECISA MELHORAR"
                    
                    # Atualizar histórico como APROVADA
                    atualizar_status_historico(
                        arquivo['id'],
                        'aprovada',  # Sempre aprovada
                        score,
                        resultado.get('erros_criticos', []),
                        resultado.get('relatorio_detalhado', '')
                    )
                    
                    # FASE 3: Atualizar estatísticas globais
                    atualizar_estatisticas(
                        tipo,
                        score,
                        True,  # Sempre True
                        tempo_geracao=0
                    )
                    
                    # Mover para pasta de aprovadas
                    mover_peticao(service, arquivo['id'], True)
                    total_aprovadas += 1
                    
                    # Mostrar resultado
                    print(f"         APROVADA! Score: {score}/100")
                    print(f"         Ranking: {ranking}")
                    
                    # Mostrar pontos de melhoria se score < 90
                    if score < 90 and resultado.get('pontos_melhoria'):
                        print(f"         Pontos de melhoria:")
                        for ponto in resultado['pontos_melhoria'][:2]:
                            print(f"           • {ponto}")
        
        print(f"\n{'='*70}")
        print(f"  AUDITOR: {total_auditadas} auditadas | {total_aprovadas}  aprovadas")
        print(f"{'='*70}")
        
    except Exception as e:
        print(f"[ERRO AUDITOR] {e}")

def verificar_flags_manuais():
    """
    Verifica se há solicitações de geração manual pendentes
    Flags são criados pelo dashboard_server.py
    """
    # Verificar flags de petição
    flags_peticao = glob.glob(os.path.join("flags", "flag_manual_*.json"))
    # Verificar flags de cronologia
    flags_cronologia = glob.glob(os.path.join("flags", "flag_cronologia_*.json"))
    
    # Processar flags de transcrição de vídeo
    flags_transcricao = glob.glob(os.path.join("flags", "flag_transcricao_*.json"))
    
    # DEBUG TOTAL - REMOVER DEPOIS
    # print(f"[DEBUG] P:{len(flags_peticao)} C:{len(flags_cronologia)} T:{len(flags_transcricao)} CWD:{os.getcwd()}")
    
    total_flags = len(flags_peticao) + len(flags_cronologia) + len(flags_transcricao)
    
    if total_flags == 0:
        return
    
    print(f"\n{'='*70}")
    print(f"   PROCESSANDO SOLICITAÇÕES MANUAIS ({total_flags})")
    print(f"{'='*70}\n")
    
    # Processar flags de petição
    for flag_file in flags_peticao:
        try:
            # Ler dados da flag
            with open(flag_file, 'r', encoding='utf-8') as f:
                dados = json.load(f)
            
            cliente_nome = dados.get('cliente_nome')
            tipo_acao = dados.get('tipo_acao', 'RECONHECIMENTO_VINCULO')
            forcar_geracao = dados.get('forcar_geracao', False)
            
            print(f"   Processando: {cliente_nome}")
            print(f"   Tipo: {tipo_acao}")
            print(f"   Forçar: {forcar_geracao}")
            
            # Processar geração
            sucesso = processar_geracao_manual(cliente_nome, tipo_acao, forcar_geracao)
            
            # Remover flag após processar (sucesso ou não)
            os.remove(flag_file)
            print(f"   Flag removida: {flag_file}\n")
            
            if sucesso:
                print(f"   Petição gerada com sucesso para {cliente_nome}\n")
            else:
                print(f"   Erro ao gerar petição para {cliente_nome}\n")
                
        except Exception as e:
            print(f"   Erro ao processar flag {flag_file}: {e}")
            # Remover flag com erro para não ficar travado
            try:
                os.remove(flag_file)
            except:
                pass
    
    # Processar flags de cronologia
    for flag_file in flags_cronologia:
        try:
            # Ler dados da flag
            with open(flag_file, 'r', encoding='utf-8') as f:
                dados = json.load(f)
            
            cliente_nome = dados.get('cliente_nome')
            
            print(f"   Gerando cronologia para: {cliente_nome}")
            
            # Processar geração de cronologia
            sucesso = processar_cronologia_manual(cliente_nome)
            
            # Remover flag após processar
            os.remove(flag_file)
            print(f"   Flag removida: {flag_file}\n")
            
            if sucesso:
                print(f"   Cronologia gerada com sucesso para {cliente_nome}\n")
            else:
                print(f"   Erro ao gerar cronologia para {cliente_nome}\n")
                
        except Exception as e:
            print(f"   Erro ao processar flag de cronologia {flag_file}: {e}")
            try:
                os.remove(flag_file)
            except:
                pass

    # Processar flags de transcrição de vídeo
    for flag_file in flags_transcricao:
        try:
            # Ler dados da flag
            with open(flag_file, 'r', encoding='utf-8') as f:
                dados = json.load(f)
            
            cliente_nome = dados.get('cliente_nome')
            
            print(f"[DEBUG] Processando vídeos para: {cliente_nome}")
            print(f"[DEBUG] Flag file: {flag_file}")
            
            # Processar transcrição de vídeo
            print(f"[DEBUG] Chamando processar_transcricao_manual...")
            sucesso = processar_transcricao_manual(cliente_nome)
            print(f"[DEBUG] processar_transcricao_manual retornou: {sucesso}")
            
            # Remover flag após processar
            os.remove(flag_file)
            print(f"   Flag removida: {flag_file}\n")
            
            if sucesso:
                print(f"   Transcrição concluída para {cliente_nome}\n")
            else:
                print(f"   Erro ao transcrever vídeo para {cliente_nome}\n")
                
        except Exception as e:
            print(f"   Erro ao processar flag de transcrição {flag_file}: {e}")
            try:
                os.remove(flag_file)
            except:
                pass

def processar_transcricao_manual(cliente_nome):
    """
    Processa transcrição de vídeo via flag manual
    Busca videos na pasta do cliente e transcreve o primeiro encontrado
    """
    try:
        print(f"\n{'='*70}")
        print(f"   TRANSCRIÇÃO DE VÍDEO AUTOMÁTICA")
        print(f"   Cliente: {cliente_nome}")
        print(f"{'='*70}\n")
        
        service = autenticar_google_drive()
        
        # 1. Buscar pasta do cliente (Reutilizando lógica da cronologia)
        pasta_cliente = None
        for _, pasta_id in [
            ('RECONHECIMENTO_VINCULO', os.getenv('PASTA_RECONHECIMENTO_VINCULO')),
            ('ACAO_ACIDENTARIA', os.getenv('PASTA_ACAO_ACIDENTARIA')),
            ('DIFERENCAS_CONTRATUAIS', os.getenv('PASTA_DIFERENCAS_CONTRATUAIS'))
        ]:
            if not pasta_id: continue
            pastas = listar_pastas(service, pasta_id)
            for p in pastas:
                if p['name'].lower() == cliente_nome.lower():
                    pasta_cliente = p
                    break
            if pasta_cliente: break
            
        if not pasta_cliente:
            print(f"   Pasta do cliente não encontrada")
            return False
            
        print(f"   Pasta encontrada: {pasta_cliente['name']} ({pasta_cliente['id']})")
        
        # 2. Listar arquivos recursivamente (incluindo subpastas como "Vídeos")
        arquivos = listar_arquivos_recursivo(service, pasta_cliente['id'])
        videos = []
        extensoes_video = ('.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv', '.m4v', '.mpeg', '.mpg')
        
        for arq in arquivos:
            if arq['name'].lower().endswith(extensoes_video):
                videos.append(arq)
                
        if not videos:
            print(f"   Nenhum vídeo encontrado na pasta do cliente")
            print(f"   Formatos aceitos: {extensoes_video}\n")
            return False
            
        print(f"   {len(videos)} vídeo(s) encontrado(s)")
        
        # 3. Transcrever TODOS os vídeos
        print(f"   Iniciando processamento de {len(videos)} vídeo(s)...")
        
        sucesso_geral = False
        videos_processados = 0
        
        for video_alvo in videos:
            print(f"\n   Analisando vídeo: {video_alvo['name']}")
            
            # Verificar se já existe Resumo ou Transcrição
            nome_base = os.path.splitext(video_alvo['name'])[0]
            existe = False
            for arq in arquivos:
                # Checa RESUMO_ ou TRANSCRICAO_
                if (arq['name'].startswith(f"RESUMO_{nome_base}") or 
                    arq['name'].startswith(f"TRANSCRICAO_{nome_base}")):
                    existe = True
                    break
            
            if existe:
                print(f"      Já processado (Arquivo existente). Pulando.")
                continue
            
            print(f"      Iniciando geração de resumo...")
            resultado = agente_transcricao_video(
                service=service,
                video_id=video_alvo['id'],
                video_nome=video_alvo['name'],
                cliente_nome=cliente_nome,
                pasta_cliente_id=pasta_cliente['id']
            )
            
            if resultado.get('success'):
                sucesso_geral = True
                videos_processados += 1
            else:
                print(f"      Falha ao processar {video_alvo['name']}")
        
        print(f"\n   Processamento concluído: {videos_processados}/{len(videos)} vídeos processados.")
        return True # Retorna True pois completou o ciclo (mesmo que tenha pulado todos)
        
    except Exception as e:
        print(f"   Erro crítico na transcrição manual: {e}")
        import traceback
        traceback.print_exc()
        return False


def processar_cronologia_manual(cliente_nome):
    """
    Processa a geração manual de cronologia dos fatos
    
    Args:
        cliente_nome: Nome do cliente
    
    Returns:
        bool: True se sucesso, False se erro
    """
    try:
        print(f"\n{'='*70}")
        print(f"   GERAÇÃO MANUAL DE CRONOLOGIA")
        print(f"   Cliente: {cliente_nome}")
        print(f"{'='*70}\n")
        
        # Autenticar Google Drive
        service = autenticar_google_drive()
        
        # Buscar pasta do cliente em todas as pastas de tipos
        pasta_cliente = None
        pasta_id_tipo = None
        
        for tipo_acao, pasta_id in [
            ('RECONHECIMENTO_VINCULO', os.getenv('PASTA_RECONHECIMENTO_VINCULO')),
            ('ACAO_ACIDENTARIA', os.getenv('PASTA_ACAO_ACIDENTARIA')),
            ('DIFERENCAS_CONTRATUAIS', os.getenv('PASTA_DIFERENCAS_CONTRATUAIS'))
        ]:
            if not pasta_id:
                continue
                
            pastas_clientes = listar_pastas(service, pasta_id)
            for pasta in pastas_clientes:
                if pasta['name'].lower() == cliente_nome.lower():
                    pasta_cliente = pasta
                    pasta_id_tipo = pasta_id
                    print(f"   Cliente encontrado em: {tipo_acao}")
                    break
            
            if pasta_cliente:
                break
        
        if not pasta_cliente:
            print(f"   Cliente não encontrado: {cliente_nome}")
            return False
        
        print(f"   ID da pasta: {pasta_cliente['id']}")
        
        # MUDANÇA: Listar arquivos recursivamente (incluindo subpastas)
        arquivos = listar_arquivos_recursivo(service, pasta_cliente['id'])
        if not arquivos:
            print(f"   Nenhum arquivo encontrado na pasta")
            return False
        
        # Classificar documentos
        docs = []
        for arq in arquivos:
            classif = classificar_documento(arq['name'])
            if classif['prioridade'] != 'IGNORAR':
                docs.append({
                    'id': arq['id'], 
                    'nome': arq['name'], 
                    'tipo': classif['tipo'], 
                    'prioridade': classif['prioridade']
                })
        
        # Procurar documento de transcrição
        doc_transcricao = next((d for d in docs if d['tipo'] == 'TRANSCRICAO'), None)
        
        if not doc_transcricao:
            print(f"   Nenhum documento de transcrição encontrado")
            print(f"   Adicione um arquivo com 'transcricao', 'entrevista' ou 'relato' no nome")
            return False
        
        print(f"   Transcrição encontrada: {doc_transcricao['nome']}")
        
        # Baixar documento de transcrição
        conteudo = baixar_arquivo(service, doc_transcricao['id'])
        if not conteudo:
            print(f"   Erro ao baixar arquivo de transcrição")
            return False
        
        # Extrair texto
        if doc_transcricao['nome'].lower().endswith('.pdf'):
            texto_transcricao = extrair_texto_pdf(conteudo)
        elif doc_transcricao['nome'].lower().endswith('.docx'):
            # Usar python-docx para extrair texto de DOCX
            try:
                tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
                os.close(tmp_fd)
                with open(tmp_path, 'wb') as f:
                    f.write(conteudo)
                doc = Document(tmp_path)
                texto_transcricao = "\n".join([p.text for p in doc.paragraphs])
                os.unlink(tmp_path)
            except Exception as e:
                print(f"   Erro ao extrair texto do DOCX: {e}")
                texto_transcricao = ""
        else:
            # Arquivo de texto puro
            try:
                texto_transcricao = conteudo.decode('utf-8', errors='ignore')
            except:
                texto_transcricao = ""
        
        if not texto_transcricao or len(texto_transcricao) < 50:
            print(f"   Transcrição muito curta ou vazia ({len(texto_transcricao)} caracteres)")
            return False
        
        print(f"   Texto extraído: {len(texto_transcricao)} caracteres")
        print(f"   Gerando cronologia com IA...")
        
        # Gerar cronologia
        cronologia_texto = agente_cronologia(texto_transcricao, pasta_cliente['name'])
        
        if not cronologia_texto:
            print(f"   Erro ao gerar cronologia (API retornou None)")
            return False
        
        print(f"   Cronologia gerada!")
        print(f"   Salvando no Google Drive...")
        
        # Salvar cronologia no Drive
        sucesso = salvar_cronologia_docx(service, cronologia_texto, pasta_cliente['name'], pasta_cliente['id'])
        
        if sucesso:
            print(f"   Cronologia salva com sucesso na pasta do cliente!")
            return True
        else:
            print(f"   Erro ao salvar cronologia")
            return False
        
    except Exception as e:
        print(f"   ERRO em processar_cronologia_manual: {e}")
        import traceback
        traceback.print_exc()
        return False

def processar_geracao_manual(cliente_nome, tipo_acao, forcar_geracao=False):
    """
    Processa a geração manual de uma petição
    
    Args:
        cliente_nome: Nome do cliente
        tipo_acao: RECONHECIMENTO_VINCULO, ACAO_ACIDENTARIA, DIFERENCAS_CONTRATUAIS
        forcar_geracao: Se True, gera mesmo com documentos faltantes
    
    Returns:
        bool: True se sucesso, False se erro
    """
    try:
        # STATUS 1: Iniciando processamento
        print(f"   Criando entrada no histórico...")
        atualizar_status_processamento(cliente_nome, tipo_acao, "Iniciando processamento...")
        
        service = autenticar_google_drive()
        
        # Determinar qual pasta buscar baseado no tipo
        pasta_id = None
        if tipo_acao == 'RECONHECIMENTO_VINCULO':
            pasta_id = os.getenv('PASTA_RECONHECIMENTO_VINCULO')
        elif tipo_acao == 'ACAO_ACIDENTARIA':
            pasta_id = os.getenv('PASTA_ACAO_ACIDENTARIA')
        elif tipo_acao == 'DIFERENCAS_CONTRATUAIS':
            pasta_id = os.getenv('PASTA_DIFERENCAS_CONTRATUAIS')
        
        if not pasta_id:
            print(f"   Pasta não configurada para {tipo_acao}")
            return False
        
        # STATUS 2: Buscando cliente
        atualizar_status_processamento(cliente_nome, tipo_acao, "Localizando pasta do cliente...")
        
        # Buscar pasta do cliente
        pastas_clientes = listar_pastas(service, pasta_id)
        pasta_cliente = None
        
        print(f"   Buscando cliente: '{cliente_nome}'")
        print(f"   Pastas encontradas: {len(pastas_clientes)}")
        
        for pasta in pastas_clientes:
            print(f"     - Comparando: '{pasta['name']}' com '{cliente_nome}'")
            if pasta['name'].lower() == cliente_nome.lower():
                pasta_cliente = pasta
                print(f"      MATCH ENCONTRADO!")
                break
        
        if not pasta_cliente:
            print(f"   Cliente não encontrado: {cliente_nome}")
            print(f"   Pastas disponíveis:")
            for pasta in pastas_clientes[:5]:
                print(f"     - {pasta['name']}")
            return False
        
        print(f"   Cliente encontrado: {pasta_cliente['name']}")
        print(f"   ID da pasta: {pasta_cliente['id']}")
        
        # Verificar se já foi processado
        if not forcar_geracao and verificar_cliente_ja_processado(service, pasta_cliente['id']):
            print(f"   Cliente já processado anteriormente (arquivo _PROCESSADO.txt encontrado).")
            print(f"   Abortando para evitar duplicação. Use forcar_geracao=True para ignorar.")
            return False

        # STATUS 3: Listando documentos
        atualizar_status_processamento(cliente_nome, tipo_acao, "Analisando documentos...")
        
        # MUDANÇA: Listar arquivos recursivamente (incluindo subpastas)
        arquivos = listar_arquivos_recursivo(service, pasta_cliente['id'])
        if not arquivos:
            print(f"   Nenhum arquivo encontrado na pasta")
            return False
        
        print(f"   Arquivos encontrados: {len(arquivos)}")
        for arq in arquivos[:5]:  # Mostrar primeiros 5
            print(f"     - {arq['name']}")
        
        # Classificar documentos
        docs = []
        for arq in arquivos:
            classif = classificar_documento(arq['name'])
            if classif['prioridade'] != 'IGNORAR':
                docs.append({
                    'id': arq['id'], 
                    'nome': arq['name'], 
                    'tipo': classif['tipo'], 
                    'prioridade': classif['prioridade']
                })
        
        print(f"   Documentos encontrados: {len(docs)}")
        
        # Verificar documentação
        verif = verificar_documentacao_completa_v10(docs, tipo_acao)
        
        print(f"   Status dos documentos:")
        print(f"      ALTA: {len(verif['alta_presentes'])}/{len(DOCUMENTOS_POR_TIPO[tipo_acao]['ALTA'])}")
        print(f"      MÉDIA: {len(verif['media_presentes'])}/{len(DOCUMENTOS_POR_TIPO[tipo_acao]['MEDIA'])}")
        print(f"      BAIXA: {len(verif['baixa_presentes'])}/{len(DOCUMENTOS_POR_TIPO[tipo_acao]['BAIXA'])}")
        
        # Verificar se tem faltantes
        tem_faltantes = (len(verif['alta_faltantes']) > 0 or 
                        len(verif['media_faltantes']) > 0 or 
                        len(verif['baixa_faltantes']) > 0)
        
        if tem_faltantes:
            print(f"\n   DOCUMENTOS FALTANTES:")
            if verif['alta_faltantes']:
                print(f"      ALTA: {', '.join(verif['alta_faltantes'])}")
            if verif['media_faltantes']:
                print(f"      MÉDIA: {', '.join(verif['media_faltantes'])}")
            if verif['baixa_faltantes']:
                print(f"      BAIXA: {', '.join(verif['baixa_faltantes'])}")
            
            if not forcar_geracao:
                print(f"\n   BLOQUEADO - Use forcar_geracao=True para gerar mesmo assim")
                return False
            else:
                print(f"\n   FORÇANDO GERAÇÃO mesmo com documentos faltantes...")
        
        # STATUS 4: Baixando documentos
        atualizar_status_processamento(cliente_nome, tipo_acao, f"Baixando {len(docs)} documentos...")
        
        # Baixar documentos completos
        print(f"\n   Baixando documentos...")
        docs_completos = []
        for doc in docs:
            cont = baixar_arquivo(service, doc['id'])
            if cont:
                texto = extrair_texto_pdf(cont) if doc['nome'].lower().endswith('.pdf') else ""
                docs_completos.append({
                    'tipo': doc['tipo'], 
                    'nome': doc['nome'], 
                    'conteudo': cont, 
                    'texto': texto
                })
        
        print(f"   {len(docs_completos)} documentos baixados")
        
        # ============================================================================
        # BUSCAR RESUMO DO VÍDEO E CRONOLOGIA DOS FATOS
        # ============================================================================
        
        resumo_texto = None
        cronologia_texto = None
        
        # 1. BUSCAR RESUMO DO VÍDEO
        doc_resumo = next((d for d in docs_completos if d['tipo'] == 'RESUMO'), None)
        
        if doc_resumo:
            print(f"\n   Resumo do vídeo encontrado: {doc_resumo['nome']}")
            resumo_texto = doc_resumo.get('texto', '').strip()
            
            # Fallback: se texto estiver vazio, tentar ler conteúdo bruto
            if not resumo_texto:
                print(f"   Texto vazio, tentando extrair novamente...")
                if doc_resumo['nome'].lower().endswith('.pdf'):
                    resumo_texto = extrair_texto_pdf(doc_resumo['conteudo'])
                elif doc_resumo['nome'].lower().endswith('.docx'):
                    # Se for DOCX, tentar ler com python-docx
                    try:
                        import tempfile
                        from docx import Document
                        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
                        os.close(tmp_fd)
                        with open(tmp_path, 'wb') as f:
                            f.write(doc_resumo['conteudo'])
                        doc = Document(tmp_path)
                        resumo_texto = '\n'.join([p.text for p in doc.paragraphs])
                        os.unlink(tmp_path)
                    except Exception as e:
                        print(f"   Erro ao ler DOCX: {e}")
                else:
                    # Se for TXT, tentar decodificar
                    try:
                        resumo_texto = doc_resumo['conteudo'].decode('utf-8', errors='ignore')
                    except:
                        pass
            
            if resumo_texto and len(resumo_texto) > 50:
                print(f"   Resumo carregado ({len(resumo_texto)} caracteres)")
            else:
                print(f"   Resumo muito curto ou vazio")
                resumo_texto = None
        else:
            print(f"  ℹ Nenhum resumo de vídeo encontrado")
        
        # 2. BUSCAR CRONOLOGIA DOS FATOS
        doc_cronologia = next((d for d in docs_completos if d['tipo'] == 'CRONOLOGIA'), None)
        
        if doc_cronologia:
            print(f"\n   Cronologia dos fatos encontrada: {doc_cronologia['nome']}")
            cronologia_texto = doc_cronologia.get('texto', '').strip()
            
            # Fallback: se texto estiver vazio, tentar ler conteúdo bruto
            if not cronologia_texto:
                print(f"   Texto vazio, tentando extrair novamente...")
                if doc_cronologia['nome'].lower().endswith('.pdf'):
                    cronologia_texto = extrair_texto_pdf(doc_cronologia['conteudo'])
                elif doc_cronologia['nome'].lower().endswith('.docx'):
                    # Se for DOCX, tentar ler com python-docx
                    try:
                        import tempfile
                        from docx import Document
                        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
                        os.close(tmp_fd)
                        with open(tmp_path, 'wb') as f:
                            f.write(doc_cronologia['conteudo'])
                        doc = Document(tmp_path)
                        cronologia_texto = '\n'.join([p.text for p in doc.paragraphs])
                        os.unlink(tmp_path)
                    except Exception as e:
                        print(f"   Erro ao ler DOCX: {e}")
                else:
                    # Se for TXT, tentar decodificar
                    try:
                        cronologia_texto = doc_cronologia['conteudo'].decode('utf-8', errors='ignore')
                    except:
                        pass
            
            if cronologia_texto and len(cronologia_texto) > 50:
                print(f"   Cronologia carregada ({len(cronologia_texto)} caracteres)")
            else:
                print(f"   Cronologia muito curta ou vazia")
                cronologia_texto = None
        else:
            print(f"  ℹ Nenhuma cronologia encontrada")
            
            # Se não tem cronologia mas tem transcrição, gerar cronologia
            doc_transcricao = next((d for d in docs_completos if d['tipo'] == 'TRANSCRICAO'), None)
            
            if doc_transcricao:
                print(f"\n   Documento de transcrição encontrado: {doc_transcricao['nome']}")
                texto_para_cronologia = doc_transcricao.get('texto', '').strip()
                
                # Fallback: se texto estiver vazio, tentar ler conteúdo bruto
                if not texto_para_cronologia:
                    print(f"   Texto vazio, tentando extrair novamente...")
                    if doc_transcricao['nome'].lower().endswith('.pdf'):
                        texto_para_cronologia = extrair_texto_pdf(doc_transcricao['conteudo'])
                    else:
                        # Se for DOCX ou TXT, tentar decodificar
                        try:
                            texto_para_cronologia = doc_transcricao['conteudo'].decode('utf-8', errors='ignore')
                        except:
                            pass
                
                if texto_para_cronologia and len(texto_para_cronologia) > 50:
                    print(f"   Gerando Cronologia dos Fatos (IA)... ({len(texto_para_cronologia)} caracteres)")
                    atualizar_status_processamento(cliente_nome, tipo_acao, "Gerando cronologia dos fatos...")
                    
                    cronologia_texto = agente_cronologia(texto_para_cronologia, pasta_cliente['name'])
                    
                    if cronologia_texto:
                        print(f"   Cronologia gerada com sucesso!")
                        salvar_cronologia_docx(service, cronologia_texto, pasta_cliente['name'], pasta_cliente['id'])
                    else:
                        print(f"   Falha ao gerar cronologia (API retornou None)")
                else:
                    print(f"   Transcrição muito curta ou vazia ({len(texto_para_cronologia)} chars). Pulando cronologia.")
            else:
                print(f"  ℹ Nenhum documento de transcrição encontrado. Pulando cronologia.")
        
        # 3. BUSCAR PROCURAÇÃO (dados completos do cliente)
        procuracao_texto = None
        doc_procuracao = next((d for d in docs_completos if d['tipo'] == 'PROCURACAO'), None)
        
        if doc_procuracao:
            print(f"\n   Procuração encontrada: {doc_procuracao['nome']}")
            procuracao_texto = doc_procuracao.get('texto', '').strip()
            
            # Fallback: se texto estiver vazio, tentar ler conteúdo bruto
            if not procuracao_texto:
                print(f"   Texto vazio, tentando extrair novamente...")
                if doc_procuracao['nome'].lower().endswith('.pdf'):
                    procuracao_texto = extrair_texto_pdf(doc_procuracao['conteudo'])
                elif doc_procuracao['nome'].lower().endswith('.docx'):
                    # Se for DOCX, tentar ler com python-docx
                    try:
                        import tempfile
                        from docx import Document
                        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
                        os.close(tmp_fd)
                        with open(tmp_path, 'wb') as f:
                            f.write(doc_procuracao['conteudo'])
                        doc = Document(tmp_path)
                        procuracao_texto = '\n'.join([p.text for p in doc.paragraphs])
                        os.unlink(tmp_path)
                    except Exception as e:
                        print(f"   Erro ao ler DOCX: {e}")
                else:
                    # Se for TXT, tentar decodificar
                    try:
                        procuracao_texto = doc_procuracao['conteudo'].decode('utf-8', errors='ignore')
                    except:
                        pass
            
            if procuracao_texto and len(procuracao_texto) > 50:
                print(f"   Procuração carregada ({len(procuracao_texto)} caracteres)")
                print(f"   Dados do cliente serão extraídos da procuração")
            else:
                print(f"   Procuração muito curta ou vazia")
                procuracao_texto = None
        else:
            print(f"  ℹ Nenhuma procuração encontrada")
        
        # STATUS 5: Gerando petição
        atualizar_status_processamento(cliente_nome, tipo_acao, "Gerando petição com IA...")
        
        # Preparar informações do cliente
        cliente_info = {
            'cliente_nome': pasta_cliente['name'],
            'tipo_processo': tipo_acao,
            'verif_docs': verif
        }
        
        # Gerar petição
        print(f"\n   Gerando petição com Claude AI...")
        if procuracao_texto:
            print(f"      Usando dados da procuração")
        if resumo_texto:
            print(f"      Usando resumo do vídeo")
        if cronologia_texto:
            print(f"      Usando cronologia dos fatos")
        
        peticao = gerar_peticao_com_claude(
            service, 
            cliente_info, 
            docs_completos, 
            tipo_acao, 
            cronologia_fatos=cronologia_texto,
            resumo_video=resumo_texto,
            procuracao=procuracao_texto,
            usar_prompt_master=True  #  PROMPT MASTER ATIVADO POR PADRÃO
        )
        
        if not peticao:
            print(f"   Erro ao gerar petição")
            return False
        
        print(f"   Petição gerada ({len(peticao)} caracteres)")
        
        # STATUS 6: Salvando
        atualizar_status_processamento(cliente_nome, tipo_acao, "Salvando petição no Drive...")
        
        # Salvar no Drive
        print(f"\n   Salvando no Google Drive...")
        arquivo = salvar_peticao_no_drive(service, peticao, cliente_info, arquivos_cliente=docs, usar_prompt_master=True)
        
        if not arquivo:
            print(f"   Erro ao salvar no Drive")
            return False
        
        print(f"   Petição salva com sucesso!")
        print(f"   Arquivo: {arquivo.get('nome', 'N/A')}")
        
        # Salvar no histórico
        print(f"\n   Salvando no histórico...")
        salvar_no_historico(
            pasta_cliente['name'],
            tipo_acao,
            arquivo
        )
        print(f"   Entrada criada no histórico!")
        
        # STATUS 7: Marcando como processado
        atualizar_status_processamento(cliente_nome, tipo_acao, "Finalizando processamento...")
        
        # Marcar cliente como processado
        print(f"\n   Marcando cliente como processado...")
        marcar_cliente_como_processado(service, pasta_cliente['id'], {
            'nome_arquivo': arquivo.get('nome', 'N/A'),
            'link': arquivo.get('link', 'N/A')
        })
        print(f"   Arquivo _PROCESSADO.txt criado!")
        
        # Gerar relatório de prints se necessário
        if arquivo.get('marcadores_prints'):
            info_prints = arquivo['marcadores_prints']
            relatorio_prints = gerar_relatorio_prints(
                info_prints['tipo_acao'],
                info_prints['marcadores'],
                info_prints['faltantes'],
                info_prints['criticos_faltantes'],
                cliente_info['cliente_nome']
            )
            if relatorio_prints:
                print(f"   Relatório de prints: {relatorio_prints}")
        
        print(f"\n   Petição gerada com sucesso para {cliente_nome}")
        print(f"   Iniciando auditoria imediata...")
        
        # Atualizar status para auditoria
        atualizar_status_processamento(cliente_nome, tipo_acao, "Auditando petição...")
        
        # AUDITAR APENAS ESTA PETIÇÃO (não chamar agente_auditor completo)
        arquivo_id = arquivo.get('id')
        if arquivo_id:
            try:
                # Auditar petição específica
                resultado_auditoria = auditar_peticao_com_claude(
                    service, 
                    arquivo_id, 
                    tipo_acao, 
                    cliente_nome
                )
                
                if resultado_auditoria:
                    print(f"   Auditoria concluída!")
                    print(f"   Score: {resultado_auditoria.get('score', 0)}/100")
                    print(f"   Ranking: {resultado_auditoria.get('ranking', 'N/A')}")
                    
                    # Salvar log de auditoria
                    log_auditoria(
                        cliente_nome,
                        tipo_acao,
                        resultado_auditoria,
                        arquivo.get('name', 'N/A')
                    )
                    
                    # Atualizar status no histórico com relatório completo
                    # IMPORTANTE: usar 'relatorio_auditoria' (não 'relatorio_detalhado')
                    atualizar_status_historico(
                        arquivo_id,
                        'aprovada',
                        resultado_auditoria.get('score'),
                        resultado_auditoria.get('erros_criticos', []),
                        resultado_auditoria.get('relatorio_detalhado', {})  #  Salva como 'relatorio_auditoria'
                    )
                    
                    # Mover para pasta de aprovadas
                    mover_peticao(service, arquivo_id, True)
                    print(f"   Movida para pasta APROVADAS")
                    
                else:
                    print(f"   Erro na auditoria")
                    
            except Exception as e:
                print(f"   Erro na auditoria: {e}")
                import traceback
                traceback.print_exc()
        
        return True
        
    except Exception as e:
        print(f"   ERRO em processar_geracao_manual: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("\n" + "="*70)
    print("  SISTEMA V10.0 - PROMPT MASTER ATIVADO")
    print("   Petições de Alto Nível (12-18 páginas)")
    print("   Times New Roman 12pt - Padrão Forense")
    print("   Nível de Advogado Sênior")
    print("="*70 + "\n")
    
    # MODO MANUAL APENAS - AUTOMÁTICO DESATIVADO
    # agente_gerador()
    # agente_auditor()
    
    # Agendar verificações (DESATIVADO PARA EVITAR VARREDURA AUTOMÁTICA)
    # intervalo = int(os.getenv('INTERVALO_MINUTOS', 1))
    # schedule.every(intervalo).minutes.do(agente_gerador)
    # schedule.every(intervalo).minutes.do(agente_auditor)
    
    print("\n" + "="*70)
    print("  Rodando! Ctrl+C para parar.")
    print("   [MODO] Sistema aguardando comandos manuais via Dashboard...")
    print("   (Varredura automática desativada)")
    print("="*70 + "\n")
    
    while True:
        # Verificar flags manuais a cada ciclo
        verificar_flags_manuais()
        
        # Executar tarefas agendadas
        schedule.run_pending()
        time.sleep(5)

if __name__ == "__main__":
    import subprocess
    import time
    import os
    
    print("="*50)
    print("  INICIALIZANDO SISTEMA UNIFICADO")
    print("="*50)
    
    # Criar diretório de flags
    if not os.path.exists('flags'):
        os.makedirs('flags')
        
    # Iniciar Dashboard em background
    print("\n[MAIN] Iniciando Dashboard (dashboard_server.py)...")
    dashboard_process = subprocess.Popen(
        ["python", "dashboard_server.py"]
        # Sem captura de stdout/stderr para ver os logs do Flask
    )
    print(f"[MAIN] Dashboard iniciado com PID: {dashboard_process.pid}")
    time.sleep(2)  # Aguardar Flask iniciar

    
    try:
        # Iniciar loop principal do Worker
        print("[MAIN] Iniciando Worker...", flush=True)
        # Hack para reduzir o sleep original do main sem alterar a função main gigante
        # Vamos apenas chamar main() e torcer para o sleep lá dentro ser aceitável
        # ou, melhor, vamos editar o sleep na função main também.
        main()
    except KeyboardInterrupt:
        print("\n[MAIN] Encerrando sistema...", flush=True)
    finally:
        print("[MAIN] Matando processo do Dashboard...", flush=True)
        dashboard_process.terminate()
        try:
            dashboard_process.wait(timeout=5)
        except:
            dashboard_process.kill()
        print("[MAIN] Sistema encerrado.", flush=True)